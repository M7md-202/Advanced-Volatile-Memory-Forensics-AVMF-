from __future__ import annotations

import base64
import hashlib
import hmac
import json
import os
import re
import shutil
import time
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

# Optional: Word reading/export
try:
    from docx import Document  # python-docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception:
    Document = None
    Table = None
    Paragraph = None
    OxmlElement = None
    qn = None

# Optional: cookie persistence
try:
    import extra_streamlit_components as stx
except Exception:
    stx = None


# =============================================================================
# Basic config / paths
# =============================================================================

st.set_page_config(page_title="AVMF Memory Forensics Dashboard", layout="wide")

BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
YARA_DIR = BASE_DIR / "yara_rules"

# New stable storage (does NOT depend on scenario name)
YARA_HITS_DIR = DATA_DIR / "yara_hits"       # per-image folders: img_<id>/
PLAYBOOKS_DIR = DATA_DIR / "playbooks"       # per-image: img_<id>.docx

DATA_DIR.mkdir(exist_ok=True)
YARA_DIR.mkdir(exist_ok=True)
YARA_HITS_DIR.mkdir(exist_ok=True)
PLAYBOOKS_DIR.mkdir(exist_ok=True)

USERS_JSON = DATA_DIR / "users.json"
REPORT_TEMPLATE_TXT = DATA_DIR / "report_template.txt"
IMAGES_JSON = DATA_DIR / "images.json"


# =============================================================================
# Table registry (single source of truth for: dashboard tabs + upload sections)
# =============================================================================

TABLE_REGISTRY: Dict[str, Dict[str, Any]] = {
    "processes": {
        "label": "Processes",
        "recommended_cols": ["pid", "ppid", "name", "user"],
        "filename_suffix": "processes",
        "group": "procnet",
        "icon": "📊",
    },
    "network": {
        "label": "Network Connections",
        "recommended_cols": ["pid", "process", "local_ip", "local_port", "remote_ip", "remote_port", "state"],
        "filename_suffix": "network",
        "group": "procnet",
        "icon": "🌐",
    },
    "yara": {
        "label": "YARA Hits",
        "recommended_cols": ["rule", "pid", "process", "address", "note"],
        "filename_suffix": "yara",  # legacy single-file
        "group": "yara",
        "icon": "🧬",
    },
    "runkeys": {
        "label": "Run Keys",
        "recommended_cols": ["key", "name", "data"],
        "filename_suffix": "runkeys",
        "group": "persistence",
        "icon": "🧷",
    },
    "runonce": {
        "label": "RunOnce",
        "recommended_cols": ["key", "name", "data"],
        "filename_suffix": "runonce",
        "group": "persistence",
        "icon": "🧷",
    },
    "cmdline": {
        "label": "Command Line",
        "recommended_cols": ["pid", "process", "cmdline"],
        "filename_suffix": "cmdline",
        "group": "cmdline",
        "icon": "🧾",
    },
    "sessions": {
        "label": "Sessions",
        "recommended_cols": ["session_id", "pid", "process", "user", "source", "destination", "protocol", "note"],
        "filename_suffix": "sessions",
        "group": "sessions",
        "icon": "🧑‍💻",
    },
    "logons": {
        "label": "Logon Events",
        "recommended_cols": ["time", "event_id", "user", "logon_type", "source_ip"],
        "filename_suffix": "logons",
        "group": "logons",
        "icon": "🔐",
    },
}

GROUP_ORDER: List[str] = ["procnet", "yara", "persistence", "cmdline", "sessions", "logons"]
GROUP_LABELS: Dict[str, str] = {
    "procnet": "📊 Processes & Network",
    "yara": "🧬 YARA",
    "persistence": "🧷 Persistence",
    "cmdline": "🧾 Cmdline",
    "sessions": "🧑‍💻 Sessions",
    "logons": "🔐 Logons & Identity",
}

DEFAULT_TABLE_VISIBILITY: Dict[str, bool] = {
    "processes": True,
    "network": True,
    "yara": True,
    "runkeys": True,
    "runonce": True,
    "cmdline": True,
    "sessions": True,
    "logons": True,
}


# =============================================================================
# Theming
# =============================================================================

def apply_theme(theme: str = "system", wide_mode: bool = True):
    if wide_mode:
        container_css = """
        .block-container { padding-top: 2rem; padding-bottom: 2rem; max-width: 100%; }
        """
    else:
        container_css = """
        .block-container { padding-top: 2rem; padding-bottom: 2rem; max-width: 1100px; margin: 0 auto; }
        """

    base_css = f"""
    :root {{ color-scheme: light dark; }}
    {container_css}

    [data-testid="stSidebar"] {{
        background-color: #111827;
        color: #e5e7eb;
    }}

    [data-testid="stSidebar"] button {{
        background: transparent !important;
        border: none !important;
        color: #e5e7eb !important;
        text-align: left !important;
        padding-left: 0 !important;
        padding-right: 0 !important;
    }}
    [data-testid="stSidebar"] button:hover {{
        color: #00e0ff !important;
        text-decoration: underline;
    }}

    .kpi-row {{
        display: flex;
        flex-wrap: wrap;
        gap: 1.25rem;
        margin: 0.5rem 0 1.2rem;
    }}
    .kpi-item {{
        flex: 1 1 0;
        min-width: 160px;
        padding: 0.25rem 0 0.6rem;
        border-bottom: 1px solid rgba(148, 163, 184, 0.35);
    }}
    .kpi-label {{
        font-size: 0.75rem;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        color: #9ca3af;
    }}
    .kpi-value {{
        font-size: 1.35rem;
        font-weight: 650;
        margin-top: 0.15rem;
    }}

    .image-info-row {{
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(210px, 1fr));
        gap: 0.9rem;
        margin-bottom: 1.2rem;
    }}
    .image-info-card {{
        border-radius: 12px;
        border: 1px solid rgba(148, 163, 184, 0.25);
        padding: 0.9rem 1rem;
        background: rgba(255,255,255,0.02);
    }}
    .image-info-label {{
        font-size: 0.72rem;
        letter-spacing: 0.12em;
        text-transform: uppercase;
        color: #9ca3af;
    }}
    .image-info-value {{
        margin-top: 0.25rem;
        font-size: 1.05rem;
        font-weight: 600;
        word-break: break-word;
        white-space: normal;
        color: #00e0ff;
    }}

    /* Playbook rendering */
    .play-wrap {{
        border: 1px solid rgba(148, 163, 184, 0.22);
        border-radius: 14px;
        padding: 1.25rem 1.25rem;
        background: rgba(0,0,0,0.15);
    }}
    .play-h1 {{
        font-size: 1.35rem;
        font-weight: 800;
        margin: 0.25rem 0 0.65rem 0;
        color: #e5e7eb;
    }}
    .play-h2 {{
        font-size: 1.10rem;
        font-weight: 750;
        margin: 1.05rem 0 0.4rem 0;
        color: #e5e7eb;
    }}
    .play-h3 {{
        font-size: 1.00rem;
        font-weight: 700;
        margin: 0.95rem 0 0.35rem 0;
        color: #e5e7eb;
        opacity: 0.95;
    }}
    .play-p {{
        font-size: 0.95rem;
        line-height: 1.55;
        margin: 0.15rem 0 0.35rem 0;
        color: rgba(229,231,235,0.92);
        white-space: pre-wrap;
    }}
    .play-li {{
        font-size: 0.95rem;
        line-height: 1.55;
        margin: 0.1rem 0 0.15rem 0;
        color: rgba(229,231,235,0.92);
        white-space: pre-wrap;
    }}

    table.pb-table {{
        width: 100%;
        border-collapse: collapse;
        margin: 0.75rem 0 1rem 0;
        border: 1px solid rgba(148, 163, 184, 0.22);
        border-radius: 10px;
        overflow: hidden;
        font-size: 0.92rem;
    }}
    .pb-table th, .pb-table td {{
        border: 1px solid rgba(148, 163, 184, 0.18);
        padding: 0.6rem 0.65rem;
        vertical-align: top;
        white-space: pre-wrap;
    }}
    .pb-table th {{
        background: rgba(255,255,255,0.04);
        font-weight: 700;
    }}
    """

    if theme == "light":
        theme_css = """
        .stApp { background-color: #ffffff; color: #111111; }
        h1,h2,h3,h4,h5,h6,label,p,span { color: #111111 !important; }
        """
    elif theme == "dark":
        theme_css = """
        .stApp { background-color: #0b1220; color: #e5e7eb; }
        h1,h2,h3,h4,h5,h6,label,p,span { color: #e5e7eb !important; }
        """
    else:
        theme_css = """
        .stApp { background-color: #ffffff; color: #111111; }
        h1,h2,h3,h4,h5,h6,label,p,span { color: #111111 !important; }
        @media (prefers-color-scheme: dark) {
            .stApp { background-color: #0b1220; color: #e5e7eb; }
            h1,h2,h3,h4,h5,h6,label,p,span { color: #e5e7eb !important; }
        }
        """

    st.markdown(f"<style>{base_css}\n{theme_css}</style>", unsafe_allow_html=True)


# =============================================================================
# Cookie manager (singleton)
# =============================================================================

def _get_cookie_manager():
    if stx is None:
        return None
    if "_cookie_manager" not in st.session_state:
        st.session_state["_cookie_manager"] = stx.CookieManager(key="avmf_cookie_manager_v1")
    return st.session_state["_cookie_manager"]


# =============================================================================
# Auth / RBAC
# =============================================================================

def _now() -> int:
    return int(time.time())


def _secret_key() -> str:
    try:
        if "AVMF_SECRET" in st.secrets:
            return str(st.secrets["AVMF_SECRET"])
    except Exception:
        pass
    return "CHANGE_ME_IN_SECRETS_OR_CODE_avmf_secret"


def _hash_password(password: str, salt: str) -> str:
    return hashlib.sha256((salt + password).encode("utf-8")).hexdigest()


def _make_token(username: str, ttl_seconds: int = 60 * 60 * 24 * 14) -> str:
    exp = _now() + ttl_seconds
    msg = f"{username}|{exp}".encode("utf-8")
    sig = hmac.new(_secret_key().encode("utf-8"), msg, hashlib.sha256).hexdigest()
    raw = f"{username}|{exp}|{sig}".encode("utf-8")
    return base64.urlsafe_b64encode(raw).decode("utf-8")


def _verify_token(token: str) -> Optional[str]:
    try:
        raw = base64.urlsafe_b64decode(token.encode("utf-8")).decode("utf-8")
        username, exp_s, sig = raw.split("|", 2)
        exp = int(exp_s)
        if _now() > exp:
            return None
        msg = f"{username}|{exp}".encode("utf-8")
        expected = hmac.new(_secret_key().encode("utf-8"), msg, hashlib.sha256).hexdigest()
        if not hmac.compare_digest(expected, sig):
            return None
        return username
    except Exception:
        return None


@dataclass
class UserCtx:
    username: str
    role: str
    perms: Dict[str, bool]


def _default_users_seed() -> Dict[str, Any]:
    return {
        "version": 1,
        "users": {
            "Mohamed_Mubarak": {
                "role": "admin",
                "salt": "s1",
                "password_hash": _hash_password("msmubrak_9750", "s1"),
                "perms": {
                    "view_dashboard": True,
                    "view_reports": True,
                    "download_reports": True,
                    "edit_report_template": True,
                    "view_playbook": True,
                    "download_playbook": True,
                    "manage_playbook": True,
                    "upload_data": True,
                    "view_yara_rules": True,
                    "edit_yara_rules": True,
                    "manage_users": True,
                    "edit_settings": True,
                    "manage_images": True,
                },
            },
            "Nora_AlSayegh": {
                "role": "employee",
                "salt": "s2",
                "password_hash": _hash_password("employee123", "s2"),
                "perms": {
                    "view_dashboard": True,
                    "view_reports": True,
                    "download_reports": True,
                    "edit_report_template": False,
                    "view_playbook": True,
                    "download_playbook": True,
                    "manage_playbook": False,
                    "upload_data": False,
                    "view_yara_rules": False,
                    "edit_yara_rules": False,
                    "manage_users": False,
                    "edit_settings": False,
                    "manage_images": False,
                },
            },
            "Omar_Hasan": {
                "role": "employee",
                "salt": "s3",
                "password_hash": _hash_password("employee123", "s3"),
                "perms": {
                    "view_dashboard": True,
                    "view_reports": True,
                    "download_reports": True,
                    "edit_report_template": False,
                    "view_playbook": True,
                    "download_playbook": True,
                    "manage_playbook": False,
                    "upload_data": False,
                    "view_yara_rules": False,
                    "edit_yara_rules": False,
                    "manage_users": False,
                    "edit_settings": False,
                    "manage_images": False,
                },
            },
        },
    }


def load_users_config() -> Dict[str, Any]:
    if not USERS_JSON.exists():
        USERS_JSON.write_text(json.dumps(_default_users_seed(), indent=2), encoding="utf-8")
    try:
        cfg = json.loads(USERS_JSON.read_text(encoding="utf-8"))
    except Exception:
        cfg = _default_users_seed()
        USERS_JSON.write_text(json.dumps(cfg, indent=2), encoding="utf-8")

    cfg.setdefault("version", 1)
    cfg.setdefault("users", {})
    return cfg


def save_users_config(cfg: Dict[str, Any]) -> None:
    USERS_JSON.write_text(json.dumps(cfg, indent=2), encoding="utf-8")


def authenticate(username: str, password: str, cfg: Dict[str, Any]) -> bool:
    users = cfg.get("users", {})
    if username not in users:
        return False
    u = users[username]
    salt = u.get("salt", "")
    ph = u.get("password_hash", "")
    return hmac.compare_digest(ph, _hash_password(password, salt))


def current_user(cfg: Dict[str, Any]) -> Optional[UserCtx]:
    # session first
    if st.session_state.get("authenticated") and st.session_state.get("username"):
        uname = st.session_state["username"]
        u = cfg["users"].get(uname)
        if u:
            return UserCtx(username=uname, role=u.get("role", "employee"), perms=u.get("perms", {}))
        return None

    # cookie fallback
    cm = _get_cookie_manager()
    if cm is None:
        return None

    token = cm.get("avmf_token")
    if not token:
        return None

    uname = _verify_token(token)
    if not uname:
        return None

    u = cfg["users"].get(uname)
    if not u:
        return None

    st.session_state["authenticated"] = True
    st.session_state["username"] = uname
    st.session_state["page"] = st.session_state.get("page", "Dashboard")
    return UserCtx(username=uname, role=u.get("role", "employee"), perms=u.get("perms", {}))


def logout():
    st.session_state["authenticated"] = False
    st.session_state["username"] = None
    st.session_state["page"] = "Dashboard"
    cm = _get_cookie_manager()
    if cm is not None:
        try:
            cm.delete("avmf_token")
        except Exception:
            pass
    st.rerun()


def can(uctx: UserCtx, perm: str) -> bool:
    return bool(uctx.perms.get(perm, False))


# =============================================================================
# Images (scenarios) persistence: data/images.json
# =============================================================================

def _seed_images() -> Dict[str, Any]:
    return {
        "version": 3,
        "images": [
            {
                "id": 1,
                "name": "mem_beacon.raw",
                "scenario": "C2 Beaconing",
                "os": "Windows 11",
                "acquired_at": "2025-01-10 14:32",
                "pipeline": "generic",
                "tabs": dict(DEFAULT_TABLE_VISIBILITY),
                "hide_empty": False,
            },
            {
                "id": 2,
                "name": "mem_fileless.raw",
                "scenario": "Ransomware (Locky)",
                "os": "Windows 11",
                "acquired_at": "2025-01-10 15:05",
                "pipeline": "locky",
                "tabs": dict(DEFAULT_TABLE_VISIBILITY),
                "hide_empty": False,
            },
            {
                "id": 3,
                "name": "mem_creds.raw",
                "scenario": "Credential Activity",
                "os": "Windows 11",
                "acquired_at": "2025-01-10 15:40",
                "pipeline": "generic",
                "tabs": dict(DEFAULT_TABLE_VISIBILITY),
                "hide_empty": False,
            },
        ],
    }


def _normalize_image_tabs(img: Dict[str, Any]) -> None:
    tabs = img.setdefault("tabs", {})
    if not isinstance(tabs, dict):
        tabs = {}
        img["tabs"] = tabs
    for k, v in DEFAULT_TABLE_VISIBILITY.items():
        tabs.setdefault(k, v)
    for k in list(tabs.keys()):
        if k not in TABLE_REGISTRY:
            tabs.pop(k, None)
    img.setdefault("hide_empty", False)


def load_images_config() -> Dict[str, Any]:
    if not IMAGES_JSON.exists():
        IMAGES_JSON.write_text(json.dumps(_seed_images(), indent=2), encoding="utf-8")
    try:
        cfg = json.loads(IMAGES_JSON.read_text(encoding="utf-8"))
    except Exception:
        cfg = _seed_images()
        IMAGES_JSON.write_text(json.dumps(cfg, indent=2), encoding="utf-8")

    cfg.setdefault("version", 3)
    cfg.setdefault("images", [])
    for img in cfg["images"]:
        img.setdefault("id", 0)
        img.setdefault("pipeline", "generic")
        img.setdefault("os", "Windows 11")
        img.setdefault("acquired_at", "")
        img.setdefault("scenario", "")
        _normalize_image_tabs(img)

    # ensure unique IDs, fill missing IDs
    used = set()
    max_id = 0
    for img in cfg["images"]:
        try:
            max_id = max(max_id, int(img.get("id", 0)))
        except Exception:
            pass
        used.add(int(img.get("id", 0) or 0))
    for img in cfg["images"]:
        if not img.get("id") or int(img.get("id")) <= 0 or int(img.get("id")) in (0,):
            max_id += 1
            img["id"] = max_id

    return cfg


def save_images_config(cfg: Dict[str, Any]) -> None:
    IMAGES_JSON.write_text(json.dumps(cfg, indent=2), encoding="utf-8")


def get_images_list(images_cfg: Dict[str, Any]) -> List[Dict[str, Any]]:
    imgs = list(images_cfg.get("images", []))
    imgs.sort(key=lambda x: int(x.get("id", 0)))
    return imgs


def next_image_id(images_cfg: Dict[str, Any]) -> int:
    ids = [int(i.get("id", 0)) for i in images_cfg.get("images", [])]
    return (max(ids) + 1) if ids else 1


def find_image(images_cfg: Dict[str, Any], image_name: str) -> Optional[Dict[str, Any]]:
    for img in images_cfg.get("images", []):
        if img.get("name") == image_name:
            return img
    return None


def find_image_by_id(images_cfg: Dict[str, Any], image_id: int) -> Optional[Dict[str, Any]]:
    for img in images_cfg.get("images", []):
        if int(img.get("id", 0)) == int(image_id):
            return img
    return None


# =============================================================================
# Storage helpers (stable per-image key)
# =============================================================================

def safe_filename(name: str) -> str:
    s = (name or "").strip()
    s = re.sub(r"[^\w\-. ]+", "_", s)
    s = s.replace(" ", "_")
    return s or "file"


def image_storage_key(images_cfg: Dict[str, Any], image_name: str) -> str:
    img = find_image(images_cfg, image_name)
    if img and int(img.get("id", 0)) > 0:
        return f"img_{int(img['id'])}"
    # fallback (should rarely happen)
    return f"imgname_{safe_filename(image_name)}"


def legacy_safe_prefix(image_name: str) -> str:
    return safe_filename(image_name).replace(".", "_")


def data_paths_for_image(images_cfg: Dict[str, Any], image_name: str) -> Dict[str, Path]:
    key = image_storage_key(images_cfg, image_name)
    paths: Dict[str, Path] = {}
    for tkey, meta in TABLE_REGISTRY.items():
        # YARA is handled separately (multi files). Keep legacy single-file too.
        paths[tkey] = DATA_DIR / f"{key}_{meta['filename_suffix']}.csv"
    return paths


def legacy_data_paths_for_image(image_name: str) -> Dict[str, Path]:
    safe = legacy_safe_prefix(image_name)
    paths: Dict[str, Path] = {}
    for tkey, meta in TABLE_REGISTRY.items():
        paths[tkey] = DATA_DIR / f"{safe}_{meta['filename_suffix']}.csv"
    return paths


def yara_hits_folder(images_cfg: Dict[str, Any], image_name: str) -> Path:
    key = image_storage_key(images_cfg, image_name)
    p = YARA_HITS_DIR / key
    p.mkdir(exist_ok=True)
    return p


def scenario_playbook_path(images_cfg: Dict[str, Any], image_name: str) -> Path:
    key = image_storage_key(images_cfg, image_name)
    return PLAYBOOKS_DIR / f"{key}.docx"


def migrate_legacy_name_based_files(images_cfg: Dict[str, Any], image_name: str) -> None:
    """
    If older name-based CSVs exist, copy them into stable ID-based filenames (only if new doesn't exist).
    This prevents "disappearing" data when scenario name changes.
    """
    new_paths = data_paths_for_image(images_cfg, image_name)
    old_paths = legacy_data_paths_for_image(image_name)

    for tkey, oldp in old_paths.items():
        newp = new_paths.get(tkey)
        if not newp:
            continue
        if newp.exists():
            continue
        if oldp.exists():
            try:
                shutil.copy2(oldp, newp)
            except Exception:
                pass

    # migrate old YARA rules file name-based -> keep as-is; rules editor uses per-image name
    # YARA hits multi-file storage is new; nothing to migrate here automatically.


# =============================================================================
# Data (blank if none)
# =============================================================================

def blank_generic_tables() -> Dict[str, pd.DataFrame]:
    return {
        "processes": pd.DataFrame(columns=["pid", "ppid", "name", "user"]),
        "network": pd.DataFrame(columns=["pid", "process", "local_ip", "local_port", "remote_ip", "remote_port", "state"]),
        "yara": pd.DataFrame(columns=["rule", "pid", "process", "address", "note", "source_file"]),
        "runkeys": pd.DataFrame(columns=["key", "name", "data"]),
        "runonce": pd.DataFrame(columns=["key", "name", "data"]),
        "cmdline": pd.DataFrame(columns=["pid", "process", "cmdline"]),
        "sessions": pd.DataFrame(columns=["session_id", "pid", "process", "user", "source", "destination", "protocol", "note"]),
        "logons": pd.DataFrame(columns=["time", "event_id", "user", "logon_type", "source_ip"]),
    }


def load_locky_from_csv_or_blank() -> Dict[str, pd.DataFrame]:
    base = blank_generic_tables()
    ps_path = DATA_DIR / "locky_pslist_norm.csv"
    pcap_path = DATA_DIR / "locky_pcap_norm.csv"
    yara_path = DATA_DIR / "locky_yara_hits.csv"
    run_path = DATA_DIR / "locky_run_keys_norm.csv"

    needed = [ps_path, pcap_path, yara_path, run_path]
    if not all(p.exists() for p in needed):
        return base

    def safe_read(p: Path) -> pd.DataFrame:
        try:
            return pd.read_csv(p)
        except Exception:
            return pd.DataFrame()

    base["processes"] = safe_read(ps_path)
    base["network"] = safe_read(pcap_path)
    y = safe_read(yara_path)
    if not y.empty and "source_file" not in y.columns:
        y["source_file"] = yara_path.name
    base["yara"] = y
    base["runkeys"] = safe_read(run_path)
    return base


def load_default_data_for_image(images_cfg: Dict[str, Any], image_name: str) -> Dict[str, pd.DataFrame]:
    img = find_image(images_cfg, image_name)
    if not img:
        return blank_generic_tables()
    pipeline = (img.get("pipeline") or "generic").lower().strip()
    if pipeline == "locky":
        return load_locky_from_csv_or_blank()
    return blank_generic_tables()


def _ensure_headers(df: pd.DataFrame, key: str) -> pd.DataFrame:
    cols = TABLE_REGISTRY[key]["recommended_cols"]
    # ensure yara has source_file for combined view
    if key == "yara" and "source_file" not in cols:
        cols = cols + ["source_file"]

    if df is None or not isinstance(df, pd.DataFrame):
        return pd.DataFrame(columns=cols)
    if df.empty and list(df.columns) == []:
        return pd.DataFrame(columns=cols)
    if list(df.columns) == []:
        return pd.DataFrame(columns=cols)

    # ensure required columns exist (without dropping extras)
    for c in cols:
        if c not in df.columns:
            df[c] = "" if c != "source_file" else ""
    return df


def _read_csv_fast(p: Path, max_rows: Optional[int] = None) -> pd.DataFrame:
    try:
        if max_rows is None:
            return pd.read_csv(p)
        return pd.read_csv(p, nrows=max_rows)
    except Exception:
        return pd.DataFrame()


def list_yara_hit_files(images_cfg: Dict[str, Any], image_name: str) -> List[Path]:
    folder = yara_hits_folder(images_cfg, image_name)
    files = sorted(folder.glob("*.csv"), key=lambda x: x.name.lower())
    return files


def load_yara_hits_combined(images_cfg: Dict[str, Any], image_name: str) -> Tuple[pd.DataFrame, List[Path]]:
    """
    Load and combine YARA hit CSVs stored in the per-image YARA_HITS folder.
    Adds a 'source_file' column to prevent confusion.
    """
    files = list_yara_hit_files(images_cfg, image_name)
    frames: List[pd.DataFrame] = []
    for p in files:
        df = _read_csv_fast(p)
        if df is None or df.empty:
            continue
        if "source_file" not in df.columns:
            df["source_file"] = p.name
        else:
            df["source_file"] = df["source_file"].astype(str).replace("", p.name)
        frames.append(df)

    if not frames:
        return pd.DataFrame(columns=["rule", "pid", "process", "address", "note", "source_file"]), files

    combined = pd.concat(frames, ignore_index=True, sort=False)
    if "source_file" not in combined.columns:
        combined["source_file"] = ""
    return combined, files


def load_image_data(images_cfg: Dict[str, Any], image_name: str) -> Dict[str, pd.DataFrame]:
    # migration (prevents "disappearing" data after rename)
    migrate_legacy_name_based_files(images_cfg, image_name)

    paths = data_paths_for_image(images_cfg, image_name)
    data = load_default_data_for_image(images_cfg, image_name)

    # per-image uploads override defaults for all except YARA (handled specially)
    for key, path in paths.items():
        if key == "yara":
            continue
        if path.exists():
            data[key] = _read_csv_fast(path)
        else:
            # if missing, keep default already loaded
            pass

    # YARA: combine multi-file storage + legacy single-file (if exists)
    combined_yara, _files = load_yara_hits_combined(images_cfg, image_name)
    legacy_single = paths.get("yara")
    if legacy_single and legacy_single.exists():
        df_single = _read_csv_fast(legacy_single)
        if not df_single.empty:
            if "source_file" not in df_single.columns:
                df_single["source_file"] = legacy_single.name
            combined_yara = pd.concat([combined_yara, df_single], ignore_index=True, sort=False)

    data["yara"] = combined_yara

    # ensure headers always
    for key in TABLE_REGISTRY.keys():
        data[key] = _ensure_headers(data.get(key, pd.DataFrame()), key)

    return data


# =============================================================================
# Reports (Word export with borders) + Timeline Reconstruction + Narrative
# =============================================================================

DEFAULT_REPORT_TEMPLATE = """AVMF FORENSIC REPORT
========================================================================

Incident Name : {incident_name}
Memory Image  : {image_name}
Operating OS  : {os}
Acquired At   : {acquired_at}

EXECUTIVE SUMMARY
------------------------------------------------------------------------
{executive_summary}

KEY FINDINGS (auto-derived from dashboard)
------------------------------------------------------------------------
{key_findings}

SUMMARY COUNTS
------------------------------------------------------------------------
Processes           : {proc_count}
Network Connections : {net_count}
YARA Hits           : {yara_count}
Run Keys            : {runkeys_count}
RunOnce             : {runonce_count}
Cmdline Rows        : {cmdline_count}
Sessions Rows       : {sessions_count}
Logon Events        : {logon_count}

TIMELINE RECONSTRUCTION (first {timeline_limit})
------------------------------------------------------------------------
{timeline_table}

EVIDENCE TABLES
------------------------------------------------------------------------

PROCESSES (first {proc_limit})
------------------------------------------------------------------------
{proc_table}

NETWORK (first {net_limit})
------------------------------------------------------------------------
{net_table}

YARA HITS (first {yara_limit})
------------------------------------------------------------------------
{yara_table}

PERSISTENCE (RUN KEYS) (first {runkeys_limit})
------------------------------------------------------------------------
{runkeys_table}

LOGON EVENTS (first {logon_limit})
------------------------------------------------------------------------
{logon_table}

ANALYST NOTES
------------------------------------------------------------------------
- Incident summary:
- Initial access vector:
- Execution chain:
- Scope & impacted hosts:
- Containment actions:
- Eradication actions:
- Recovery actions:
- Lessons learned:
"""


def load_report_template() -> str:
    if not REPORT_TEMPLATE_TXT.exists():
        REPORT_TEMPLATE_TXT.write_text(DEFAULT_REPORT_TEMPLATE, encoding="utf-8")
    try:
        return REPORT_TEMPLATE_TXT.read_text(encoding="utf-8")
    except Exception:
        return DEFAULT_REPORT_TEMPLATE


def save_report_template(text: str) -> None:
    REPORT_TEMPLATE_TXT.write_text(text, encoding="utf-8")


def _df_block(df: pd.DataFrame, n: int = 10) -> str:
    if df is None:
        return ""
    if df.empty:
        try:
            return df.head(0).to_string(index=False)
        except Exception:
            return ""
    return df.head(n).to_string(index=False)


def _coerce_dt_series(s: pd.Series) -> pd.Series:
    # best-effort datetime parse; produce UTC-aware series
    try:
        dt = pd.to_datetime(s, errors="coerce", utc=True)
        return dt
    except Exception:
        try:
            dt = pd.to_datetime(s.astype(str), errors="coerce", utc=True)
            return dt
        except Exception:
            return pd.Series([pd.NaT] * len(s))


def _find_col(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    cols = {c.lower(): c for c in df.columns}
    for cand in candidates:
        if cand.lower() in cols:
            return cols[cand.lower()]
    return None


# -------------------------------
# Narrative helpers (NEW)
# -------------------------------

def _try_parse_dt(value: Any) -> Optional[datetime]:
    if value is None:
        return None
    try:
        dt = pd.to_datetime(value, errors="coerce", utc=True)
        if pd.isna(dt):
            return None
        return dt.to_pydatetime() if hasattr(dt, "to_pydatetime") else dt
    except Exception:
        return None


def _format_dt(dt: Optional[datetime], tz_label: str = "UTC") -> str:
    if not dt:
        return "Unknown"
    try:
        return dt.strftime(f"%Y-%m-%d %H:%M:%S {tz_label}")
    except Exception:
        return str(dt)


def _infer_incident_start_time(data: Dict[str, pd.DataFrame], timeline_df: pd.DataFrame) -> Optional[datetime]:
    """
    Priority:
    1) Earliest process creation/start time (best indicator of 'incident took place at')
    2) Earliest timestamp anywhere in timeline reconstruction (fallback)
    """
    proc = data.get("processes", pd.DataFrame())
    if isinstance(proc, pd.DataFrame) and not proc.empty:
        # Look for likely process start/create columns (covers Volatility 'CreateTime', 'StartTime', etc.)
        candidates: List[str] = []
        for c in proc.columns:
            cl = c.lower()
            if "time" in cl and any(k in cl for k in ["create", "start", "created", "started", "launch"]):
                candidates.append(c)

        # Also include common exact variants
        exact = ["create_time", "creation_time", "start_time", "created_at", "started_at", "createtime", "starttime"]
        for e in exact:
            col = _find_col(proc, [e])
            if col and col not in candidates:
                candidates.append(col)

        best_min: Optional[datetime] = None
        for c in candidates:
            dt_series = _coerce_dt_series(proc[c])
            if dt_series is None or dt_series.empty:
                continue
            try:
                dt_min = dt_series.dropna().min()
                if pd.isna(dt_min):
                    continue
                dt_min = dt_min.to_pydatetime() if hasattr(dt_min, "to_pydatetime") else dt_min
                if best_min is None or dt_min < best_min:
                    best_min = dt_min
            except Exception:
                continue

        if best_min:
            return best_min

    # fallback: earliest timeline time_utc (string) -> parse
    if isinstance(timeline_df, pd.DataFrame) and not timeline_df.empty and "time_utc" in timeline_df.columns:
        parsed = pd.to_datetime(timeline_df["time_utc"], errors="coerce", utc=True)
        parsed = parsed.dropna()
        if not parsed.empty:
            dt_min = parsed.min()
            return dt_min.to_pydatetime() if hasattr(dt_min, "to_pydatetime") else dt_min

    return None


def _top_values(df: pd.DataFrame, col_candidates: List[str], top_n: int = 5) -> List[Tuple[str, int]]:
    if df is None or df.empty:
        return []
    col = _find_col(df, col_candidates)
    if not col:
        return []
    try:
        vc = df[col].astype(str).replace("nan", "").replace("None", "").str.strip()
        vc = vc[vc != ""]
        if vc.empty:
            return []
        out = vc.value_counts().head(top_n)
        return [(str(k), int(v)) for k, v in out.items()]
    except Exception:
        return []


def build_executive_summary_and_findings(
    img_meta: Dict[str, Any],
    data: Dict[str, pd.DataFrame],
    timeline_df: pd.DataFrame
) -> Tuple[str, str]:
    incident_name = (img_meta.get("scenario") or "").strip() or "Unspecified Incident"
    acquired_at_raw = str(img_meta.get("acquired_at", "") or "").strip()
    acquired_at_display = acquired_at_raw if acquired_at_raw else "Unknown"

    incident_start = _infer_incident_start_time(data, timeline_df)
    incident_start_display = _format_dt(incident_start, "UTC")

    proc_count = len(data.get("processes", pd.DataFrame()))
    net_count = len(data.get("network", pd.DataFrame()))
    yara_count = len(data.get("yara", pd.DataFrame()))
    runkeys_count = len(data.get("runkeys", pd.DataFrame()))
    runonce_count = len(data.get("runonce", pd.DataFrame()))
    sessions_count = len(data.get("sessions", pd.DataFrame()))
    logon_count = len(data.get("logons", pd.DataFrame()))
    cmdline_count = len(data.get("cmdline", pd.DataFrame()))

    executive = (
        f"The incident **{incident_name}** is assessed to have occurred around **{incident_start_display}**, "
        f"based on the earliest available process start/creation timestamp in the memory artifacts. "
        f"Memory acquisition was performed at **{acquired_at_display}**. "
        f"Analysis of the image identified **{proc_count} processes**, **{net_count} network entries**, "
        f"and **{yara_count} YARA hits**, alongside persistence and identity artifacts where available."
    )

    bullets: List[str] = []

    top_yara = _top_values(data.get("yara", pd.DataFrame()), ["rule", "rulename"], 5)
    if top_yara:
        bullets.append("YARA detections observed (top rules): " + ", ".join([f"{r} ({c})" for r, c in top_yara]))

    top_remote = _top_values(data.get("network", pd.DataFrame()), ["remote_ip", "raddr", "dst_ip", "remoteaddress"], 5)
    if top_remote:
        bullets.append("Most frequent remote IPs: " + ", ".join([f"{ip} ({c})" for ip, c in top_remote]))

    top_proc = _top_values(
        data.get("processes", pd.DataFrame()),
        ["name", "process", "image_file_name", "imagefilename", "imagefilename", "imagefilename", "imagefilename", "imagefilename", "ImageFileName"],
        8
    )
    if top_proc:
        bullets.append("Most frequent process names: " + ", ".join([f"{p} ({c})" for p, c in top_proc]))

    if runkeys_count or runonce_count:
        bullets.append(f"Persistence artifacts present: Run Keys={runkeys_count}, RunOnce={runonce_count}")

    if logon_count:
        top_users = _top_values(data.get("logons", pd.DataFrame()), ["user", "account", "username"], 5)
        if top_users:
            bullets.append("Users observed in logon artifacts: " + ", ".join([f"{u} ({c})" for u, c in top_users]))
        top_src = _top_values(data.get("logons", pd.DataFrame()), ["source_ip", "ip", "src_ip"], 5)
        if top_src:
            bullets.append("Source IPs in logon artifacts: " + ", ".join([f"{ip} ({c})" for ip, c in top_src]))

    if sessions_count:
        bullets.append(f"Session artifacts present: {sessions_count} row(s)")

    if cmdline_count:
        bullets.append(f"Command-line artifacts present: {cmdline_count} row(s)")

    if not bullets:
        bullets.append("No high-confidence findings were derived from the uploaded datasets for this scenario.")

    findings_text = "\n".join([f"- {b}" for b in bullets])
    return executive, findings_text


# =============================================================================
# Timeline Reconstruction
# =============================================================================

def build_timeline_df(data: Dict[str, pd.DataFrame], limit: int = 50) -> pd.DataFrame:
    """
    Only uses sources that actually include a parseable time column.
    For processes: tries to use create/start time + exit/end time.
    For others: uses the most obvious time column.
    """
    events: List[Dict[str, Any]] = []

    # helper: pid / process col
    def get_pid(df: pd.DataFrame) -> Optional[str]:
        return _find_pid_column(df)

    def get_proc(df: pd.DataFrame) -> Optional[str]:
        # Added "imagefilename" / "imagefile" variants (Volatility often uses ImageFileName)
        return _find_col(df, ["process", "name", "image_file_name", "imagefilename", "imagefile", "exe", "process_name", "ImageFileName"])

    # PROCESSES
    proc = data.get("processes", pd.DataFrame())
    if isinstance(proc, pd.DataFrame) and not proc.empty:
        pid_col = get_pid(proc)
        proc_col = get_proc(proc)

        start_col = None
        end_col = None
        # common names
        for c in proc.columns:
            cl = c.lower()
            if start_col is None and any(k in cl for k in ["create", "start", "created", "started"]) and "time" in cl:
                start_col = c
            if start_col is None and cl in ("create_time", "creation_time", "start_time", "created_time", "createtime", "starttime"):
                start_col = c
            if end_col is None and any(k in cl for k in ["exit", "end", "terminate", "terminated"]) and "time" in cl:
                end_col = c
            if end_col is None and cl in ("exit_time", "end_time", "terminated_time", "exittime", "endtime"):
                end_col = c

        if start_col:
            dt = _coerce_dt_series(proc[start_col])
            for i in range(len(proc)):
                if pd.isna(dt.iloc[i]):
                    continue
                events.append({
                    "time_utc": dt.iloc[i].to_pydatetime(),
                    "source": "processes",
                    "pid": str(proc.iloc[i][pid_col]) if pid_col else "",
                    "process": str(proc.iloc[i][proc_col]) if proc_col else "",
                    "event": "Process start",
                })
        if end_col:
            dt = _coerce_dt_series(proc[end_col])
            for i in range(len(proc)):
                if pd.isna(dt.iloc[i]):
                    continue
                events.append({
                    "time_utc": dt.iloc[i].to_pydatetime(),
                    "source": "processes",
                    "pid": str(proc.iloc[i][pid_col]) if pid_col else "",
                    "process": str(proc.iloc[i][proc_col]) if proc_col else "",
                    "event": "Process exit",
                })

    # LOGONS
    logons = data.get("logons", pd.DataFrame())
    if isinstance(logons, pd.DataFrame) and not logons.empty:
        tcol = _find_time_column(logons)
        pid_col = get_pid(logons)
        proc_col = get_proc(logons)
        if tcol:
            dt = _coerce_dt_series(logons[tcol])
            for i in range(len(logons)):
                if pd.isna(dt.iloc[i]):
                    continue
                user = str(logons.iloc[i].get(_find_col(logons, ["user"]), ""))
                eid = str(logons.iloc[i].get(_find_col(logons, ["event_id", "eventid", "id"]), ""))
                ltype = str(logons.iloc[i].get(_find_col(logons, ["logon_type", "type"]), ""))
                src = str(logons.iloc[i].get(_find_col(logons, ["source_ip", "ip", "src_ip"]), ""))
                ev = f"Logon event_id={eid} user={user} type={ltype} src={src}".strip()
                events.append({
                    "time_utc": dt.iloc[i].to_pydatetime(),
                    "source": "logons",
                    "pid": str(logons.iloc[i][pid_col]) if pid_col else "",
                    "process": str(logons.iloc[i][proc_col]) if proc_col else "",
                    "event": ev,
                })

    # NETWORK
    net = data.get("network", pd.DataFrame())
    if isinstance(net, pd.DataFrame) and not net.empty:
        tcol = _find_time_column(net)
        pid_col = get_pid(net)
        proc_col = get_proc(net)
        if tcol:
            dt = _coerce_dt_series(net[tcol])
            for i in range(len(net)):
                if pd.isna(dt.iloc[i]):
                    continue
                lip = str(net.iloc[i].get(_find_col(net, ["local_ip", "laddr", "src_ip"]), ""))
                lpt = str(net.iloc[i].get(_find_col(net, ["local_port", "lport", "src_port"]), ""))
                rip = str(net.iloc[i].get(_find_col(net, ["remote_ip", "raddr", "dst_ip"]), ""))
                rpt = str(net.iloc[i].get(_find_col(net, ["remote_port", "rport", "dst_port"]), ""))
                stt = str(net.iloc[i].get(_find_col(net, ["state", "status"]), ""))
                ev = f"Network {lip}:{lpt} -> {rip}:{rpt} state={stt}".strip()
                events.append({
                    "time_utc": dt.iloc[i].to_pydatetime(),
                    "source": "network",
                    "pid": str(net.iloc[i][pid_col]) if pid_col else "",
                    "process": str(net.iloc[i][proc_col]) if proc_col else "",
                    "event": ev,
                })

    # SESSIONS
    ses = data.get("sessions", pd.DataFrame())
    if isinstance(ses, pd.DataFrame) and not ses.empty:
        tcol = _find_time_column(ses)
        pid_col = get_pid(ses)
        proc_col = get_proc(ses)
        if tcol:
            dt = _coerce_dt_series(ses[tcol])
            for i in range(len(ses)):
                if pd.isna(dt.iloc[i]):
                    continue
                src = str(ses.iloc[i].get(_find_col(ses, ["source", "src", "src_ip"]), ""))
                dst = str(ses.iloc[i].get(_find_col(ses, ["destination", "dst", "dst_ip"]), ""))
                proto = str(ses.iloc[i].get(_find_col(ses, ["protocol", "proto"]), ""))
                sid = str(ses.iloc[i].get(_find_col(ses, ["session_id", "sessionid", "id"]), ""))
                ev = f"Session id={sid} {src} -> {dst} proto={proto}".strip()
                events.append({
                    "time_utc": dt.iloc[i].to_pydatetime(),
                    "source": "sessions",
                    "pid": str(ses.iloc[i][pid_col]) if pid_col else "",
                    "process": str(ses.iloc[i][proc_col]) if proc_col else "",
                    "event": ev,
                })

    # YARA (only if a time column exists)
    yara = data.get("yara", pd.DataFrame())
    if isinstance(yara, pd.DataFrame) and not yara.empty:
        tcol = _find_time_column(yara)
        pid_col = get_pid(yara)
        proc_col = get_proc(yara)
        if tcol:
            dt = _coerce_dt_series(yara[tcol])
            rule_col = _find_col(yara, ["rule", "rulename"])
            srcf = _find_col(yara, ["source_file", "file"])
            for i in range(len(yara)):
                if pd.isna(dt.iloc[i]):
                    continue
                rule = str(yara.iloc[i].get(rule_col, "")) if rule_col else ""
                sf = str(yara.iloc[i].get(srcf, "")) if srcf else ""
                ev = f"YARA hit rule={rule} file={sf}".strip()
                events.append({
                    "time_utc": dt.iloc[i].to_pydatetime(),
                    "source": "yara",
                    "pid": str(yara.iloc[i][pid_col]) if pid_col else "",
                    "process": str(yara.iloc[i][proc_col]) if proc_col else "",
                    "event": ev,
                })

    if not events:
        return pd.DataFrame(columns=["time_utc", "source", "pid", "process", "event"])

    df = pd.DataFrame(events)
    df = df.dropna(subset=["time_utc"])
    df = df.sort_values("time_utc", ascending=True)
    df["time_utc"] = df["time_utc"].apply(lambda x: x.strftime("%Y-%m-%d %H:%M:%S UTC") if isinstance(x, datetime) else str(x))
    return df.head(limit).reset_index(drop=True)


def render_report_from_dashboard(images_cfg: Dict[str, Any], template_text: str, image_name: str) -> Tuple[str, Dict[str, Any]]:
    img_meta = find_image(images_cfg, image_name) or {"name": image_name, "scenario": "", "os": "", "acquired_at": ""}
    data = load_image_data(images_cfg, image_name)

    proc_limit = 10
    net_limit = 10
    yara_limit = 10
    runkeys_limit = 10
    logon_limit = 10
    timeline_limit = 50

    timeline_df = build_timeline_df(data, limit=timeline_limit)

    # NEW: narrative blocks
    executive_summary, key_findings = build_executive_summary_and_findings(img_meta, data, timeline_df)

    mapping = {
        "{incident_name}": str(img_meta.get("scenario", "") or "Unspecified Incident"),

        "{image_name}": str(img_meta.get("name", "")),
        "{scenario}": str(img_meta.get("scenario", "")),
        "{os}": str(img_meta.get("os", "")),
        "{acquired_at}": str(img_meta.get("acquired_at", "")),

        "{executive_summary}": executive_summary,
        "{key_findings}": key_findings,

        "{proc_count}": str(len(data["processes"].index)),
        "{net_count}": str(len(data["network"].index)),
        "{yara_count}": str(len(data["yara"].index)),
        "{runkeys_count}": str(len(data["runkeys"].index)),
        "{runonce_count}": str(len(data["runonce"].index)),
        "{cmdline_count}": str(len(data["cmdline"].index)),
        "{sessions_count}": str(len(data["sessions"].index)),
        "{logon_count}": str(len(data["logons"].index)),

        "{proc_limit}": str(proc_limit),
        "{net_limit}": str(net_limit),
        "{yara_limit}": str(yara_limit),
        "{runkeys_limit}": str(runkeys_limit),
        "{logon_limit}": str(logon_limit),
        "{timeline_limit}": str(timeline_limit),

        "{timeline_table}": _df_block(timeline_df, timeline_limit),
        "{proc_table}": _df_block(data["processes"], proc_limit),
        "{net_table}": _df_block(data["network"], net_limit),
        "{yara_table}": _df_block(data["yara"], yara_limit),
        "{runkeys_table}": _df_block(data["runkeys"], runkeys_limit),
        "{logon_table}": _df_block(data["logons"], logon_limit),
    }

    out = template_text
    for k, v in mapping.items():
        out = out.replace(k, v)

    payload = {"meta": img_meta, "data": data, "timeline": timeline_df}
    return out, payload


def _set_table_borders_word(table):
    if OxmlElement is None or qn is None:
        return
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        tblBorders.append(element)
    tblPr.append(tblBorders)


def _docx_add_df_table(doc: Document, df: pd.DataFrame, title: str, limit: int):
    doc.add_heading(title, level=2)
    if df is None:
        df = pd.DataFrame()

    view = df.head(limit).copy() if not df.empty else df.head(0).copy()
    cols = list(view.columns)
    if not cols:
        cols = [""]

    table = doc.add_table(rows=1, cols=len(cols))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    _set_table_borders_word(table)

    hdr_cells = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr_cells[i].text = str(c)

    if not df.empty:
        for _, row in df.head(limit).iterrows():
            r = table.add_row().cells
            for i, c in enumerate(cols):
                val = row.get(c, "")
                r[i].text = "" if pd.isna(val) else str(val)


def build_report_docx_bytes(payload: Dict[str, Any]) -> Optional[bytes]:
    if Document is None:
        return None

    doc = Document()
    meta = payload["meta"]
    data: Dict[str, pd.DataFrame] = payload["data"]
    timeline: pd.DataFrame = payload.get("timeline", pd.DataFrame())

    incident_name = str(meta.get("scenario", "") or "Unspecified Incident")

    doc.add_heading("AVMF Forensic Report", level=0)
    doc.add_paragraph(f"Incident Name: {incident_name}")
    doc.add_paragraph(f"Memory Image: {meta.get('name', '')}")
    doc.add_paragraph(f"Operating OS: {meta.get('os', '')}")
    doc.add_paragraph(f"Acquired At: {meta.get('acquired_at', '')}")

    # NEW: Narrative blocks in DOCX
    executive_summary, key_findings = build_executive_summary_and_findings(meta, data, timeline)

    doc.add_paragraph("")
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(executive_summary)

    doc.add_paragraph("")
    doc.add_heading("Key Findings (auto-derived)", level=1)
    for line in (key_findings or "").splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("-"):
            p = doc.add_paragraph(line.lstrip("-").strip())
            try:
                p.style = "List Bullet"
            except Exception:
                pass
        else:
            doc.add_paragraph(line)

    doc.add_paragraph("")
    doc.add_heading("Summary Counts", level=1)
    doc.add_paragraph(f"Processes: {len(data['processes'])}")
    doc.add_paragraph(f"Network Connections: {len(data['network'])}")
    doc.add_paragraph(f"YARA Hits: {len(data['yara'])}")
    doc.add_paragraph(f"Run Keys: {len(data['runkeys'])}")
    doc.add_paragraph(f"RunOnce: {len(data['runonce'])}")
    doc.add_paragraph(f"Cmdline Rows: {len(data['cmdline'])}")
    doc.add_paragraph(f"Sessions Rows: {len(data['sessions'])}")
    doc.add_paragraph(f"Logon Events: {len(data['logons'])}")

    doc.add_paragraph("")
    _docx_add_df_table(doc, timeline, "Timeline Reconstruction (first 50)", 50)
    _docx_add_df_table(doc, data["processes"], "Processes (first 10)", 10)
    _docx_add_df_table(doc, data["network"], "Network (first 10)", 10)
    _docx_add_df_table(doc, data["yara"], "YARA Hits (first 10)", 10)
    _docx_add_df_table(doc, data["runkeys"], "Persistence Run Keys (first 10)", 10)
    _docx_add_df_table(doc, data["runonce"], "Persistence RunOnce (first 10)", 10)
    _docx_add_df_table(doc, data["cmdline"], "Command Line (first 10)", 10)
    _docx_add_df_table(doc, data["sessions"], "Sessions (first 10)", 10)
    _docx_add_df_table(doc, data["logons"], "Logon Events (first 10)", 10)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# =============================================================================
# YARA per-rule editor (kept per scenario name; independent of YARA hits CSVs)
# =============================================================================

def _yara_file_for_image(image_name: str) -> Path:
    return YARA_DIR / f"{image_name.replace('.', '_')}.yar"


def ensure_default_yara_file(image_name: str) -> Path:
    p = _yara_file_for_image(image_name)
    if not p.exists():
        p.write_text(f"// YARA rules for {image_name}\n\n", encoding="utf-8")
    return p


def split_yara_rules(text: str) -> Tuple[str, List[Tuple[str, str]]]:
    if not text:
        return "", []
    pattern = re.compile(r"(?m)^\s*rule\s+([A-Za-z0-9_]+)\b")
    matches = list(pattern.finditer(text))
    if not matches:
        return text.strip(), []

    preamble = text[:matches[0].start()].rstrip("\n")
    rules: List[Tuple[str, str]] = []
    for i, m in enumerate(matches):
        name = m.group(1)
        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        rules.append((name, text[start:end].strip("\n")))
    return preamble, rules


def merge_yara_rules(preamble: str, rules: List[Tuple[str, str]]) -> str:
    chunks: List[str] = []
    pre = (preamble or "").strip("\n")
    if pre:
        chunks.append(pre)
    for _, rt in rules:
        rt = (rt or "").strip("\n")
        if rt:
            chunks.append(rt)
    return "\n\n".join(chunks).strip() + "\n"


def parse_rule_name_from_text(rule_text: str) -> Optional[str]:
    m = re.search(r"(?m)^\s*rule\s+([A-Za-z0-9_]+)\b", rule_text or "")
    return m.group(1) if m else None


def _yara_state_keys(image_name: str) -> Dict[str, str]:
    sel_key = f"yara_rule_select__{image_name}"
    editor_key = f"yara_rule_editor__{image_name}"
    last_sel_key = f"yara_last_rule_select__{image_name}"
    pending_sel_key = f"{sel_key}__pending"
    pending_editor_key = f"{editor_key}__pending"
    return {
        "sel": sel_key,
        "editor": editor_key,
        "last": last_sel_key,
        "pending_sel": pending_sel_key,
        "pending_editor": pending_editor_key,
    }


def yara_rules_page(image_name: str, uctx: UserCtx):
    st.title("YARA Rules Library")

    if not can(uctx, "view_yara_rules"):
        st.warning("Admins only page (employees cannot access).")
        return

    rules_path = ensure_default_yara_file(image_name)
    text = rules_path.read_text(encoding="utf-8")
    preamble, rules = split_yara_rules(text)
    rule_names = [name for name, _ in rules]

    keys = _yara_state_keys(image_name)

    pending_sel = st.session_state.get(keys["pending_sel"])
    if pending_sel is not None:
        if pending_sel == "" or pending_sel in rule_names:
            st.session_state[keys["sel"]] = pending_sel
        st.session_state.pop(keys["pending_sel"], None)

    pending_editor = st.session_state.get(keys["pending_editor"])
    if pending_editor is not None:
        st.session_state[keys["editor"]] = pending_editor
        st.session_state.pop(keys["pending_editor"], None)

    st.markdown(f"**Memory image:** `{image_name}`")
    st.caption("Rules are edited per-rule (easier add/remove). Final download is one combined .yar.")

    with st.expander("Preamble (imports / globals / comments)", expanded=False):
        preamble_new = st.text_area("Preamble", value=preamble, height=150, key=f"yara_preamble_{image_name}")

    if not rule_names:
        st.info("No rules found. Add one below.")
        selected_name = ""
        selected_idx = None
        selected_text = ""
    else:
        if keys["sel"] not in st.session_state or st.session_state.get(keys["sel"]) not in rule_names:
            st.session_state[keys["sel"]] = rule_names[0]

        selected_name = st.selectbox("Select a rule", rule_names, key=keys["sel"])
        selected_idx = rule_names.index(selected_name)
        selected_text = rules[selected_idx][1]

    if st.session_state.get(keys["last"]) != selected_name:
        st.session_state[keys["last"]] = selected_name
        st.session_state[keys["editor"]] = selected_text or ""

    st.subheader("Rule Editor")
    if selected_name:
        st.markdown(f"Editing: **{selected_name}**")

    rule_text_new = st.text_area(
        "Rule text",
        value=st.session_state.get(keys["editor"], ""),
        height=420,
        key=keys["editor"],
    )

    c1, c2, c3, c4 = st.columns([1, 1, 1, 1])

    def save_all(preamble_text: str, rules_list: List[Tuple[str, str]]):
        final = merge_yara_rules(preamble_text, rules_list)
        rules_path.write_text(final, encoding="utf-8")

    with c1:
        if st.button("Save rule", type="primary", key=f"yara_save_rule_{image_name}"):
            if not selected_name or selected_idx is None:
                st.error("No rule selected.")
            else:
                header_name = (parse_rule_name_from_text(rule_text_new) or selected_name).strip()

                new_rules: List[Tuple[str, str]] = []
                for i, (nm, rt) in enumerate(rules):
                    if i != selected_idx:
                        new_rules.append((nm, rt))
                    else:
                        new_rules.append((header_name, (rule_text_new or "").strip("\n")))

                names_after = [n for n, _ in new_rules]
                if len(names_after) != len(set(names_after)):
                    st.error("Duplicate rule name detected. Pick a unique rule name.")
                else:
                    rules = new_rules
                    save_all(preamble_new, rules)
                    st.session_state[keys["pending_sel"]] = header_name
                    st.rerun()

    with c2:
        if st.button("Delete rule", key=f"yara_delete_rule_{image_name}"):
            if not selected_name or selected_idx is None:
                st.error("No rule selected.")
            else:
                rules.pop(selected_idx)
                save_all(preamble_new, rules)
                remaining_names = [n for n, _ in rules]
                st.session_state[keys["pending_sel"]] = remaining_names[0] if remaining_names else ""
                st.session_state[keys["pending_editor"]] = ""
                st.rerun()

    with c3:
        new_name = st.text_input("New rule name", value="", key=f"yara_new_rule_name_{image_name}")
        if st.button("Add new rule", key=f"yara_add_rule_{image_name}"):
            nn = (new_name or "").strip()
            if not nn:
                st.error("Enter a new rule name first.")
            elif nn in rule_names:
                st.error("Rule name already exists.")
            else:
                skeleton = f"""rule {nn}
{{
    meta:
        description = "Describe what this rule detects"
        author = "{uctx.username}"
        reference = "Optional reference"

    strings:
        $a1 = "example" nocase

    condition:
        any of them
}}"""
                rules.append((nn, skeleton.strip("\n")))
                save_all(preamble_new, rules)
                st.session_state[keys["pending_sel"]] = nn
                st.session_state[keys["pending_editor"]] = skeleton
                st.rerun()

    with c4:
        final_text = merge_yara_rules(preamble_new, rules)
        st.download_button(
            "Download .yar",
            data=final_text.encode("utf-8"),
            file_name=rules_path.name,
            mime="text/plain",
            key=f"yara_download_{image_name}",
        )


# =============================================================================
# IR Playbook (per scenario)
# =============================================================================

@dataclass
class PBPara:
    kind: str
    payload: Any


@dataclass
class PBSection:
    title: str
    level: int
    paras: List[PBPara]


def _iter_block_items(doc: Document) -> List[Union["Paragraph", "Table"]]:
    if Document is None or Paragraph is None or Table is None:
        return []
    items: List[Union[Paragraph, Table]] = []
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            items.append(Paragraph(child, doc))
        elif child.tag.endswith("}tbl"):
            items.append(Table(child, doc))
    return items


def _table_to_matrix(tbl: Table) -> List[List[str]]:
    matrix: List[List[str]] = []
    for row in tbl.rows:
        r: List[str] = []
        for cell in row.cells:
            r.append((cell.text or "").strip("\n"))
        matrix.append(r)
    matrix = [r for r in matrix if any((c or "").strip() for c in r)]
    return matrix


def _docx_to_playbook_sections(doc: Document) -> List[PBSection]:
    sections: List[PBSection] = []
    current: Optional[PBSection] = None

    def start_section(title: str, level: int):
        nonlocal current
        current = PBSection(title=title, level=level, paras=[])
        sections.append(current)

    def add_para(kind: str, payload: Any):
        nonlocal current
        if current is None:
            start_section("Playbook", 1)
        current.paras.append(PBPara(kind=kind, payload=payload))

    blocks = _iter_block_items(doc)
    for block in blocks:
        if Paragraph is not None and isinstance(block, Paragraph):
            text = (block.text or "").strip()
            if not text:
                continue
            style_name = ""
            try:
                style_name = (block.style.name or "").lower()
            except Exception:
                style_name = ""

            if "heading 1" in style_name:
                start_section(text, 1)
                continue
            if "heading 2" in style_name:
                start_section(text, 2)
                continue
            if "heading 3" in style_name:
                start_section(text, 3)
                continue

            if re.match(r"^\d+\.\s+", text):
                start_section(text, 1)
                continue
            if re.match(r"^\d+\.\d+\s+", text):
                start_section(text, 2)
                continue
            if re.match(r"^\d+\.\d+\.\d+\s+", text):
                start_section(text, 3)
                continue

            if text.startswith(("-", "•", "‣")):
                add_para("li", text.lstrip("-•‣").strip())
            else:
                add_para("p", text)

        elif Table is not None and isinstance(block, Table):
            matrix = _table_to_matrix(block)
            if matrix:
                add_para("table", matrix)

    return sections


def _escape_html(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")).replace("\n", "<br/>")


def _render_playbook_html(paras: List[PBPara]) -> str:
    chunks: List[str] = []
    chunks.append('<div class="play-wrap">')
    for para in paras:
        if para.kind in ("h1", "h2", "h3", "p", "li"):
            t = _escape_html(str(para.payload or "").strip())
            if para.kind == "h1":
                chunks.append(f'<div class="play-h1">{t}</div>')
            elif para.kind == "h2":
                chunks.append(f'<div class="play-h2">{t}</div>')
            elif para.kind == "h3":
                chunks.append(f'<div class="play-h3">{t}</div>')
            elif para.kind == "li":
                chunks.append(f'<div class="play-li">• {t}</div>')
            else:
                chunks.append(f'<div class="play-p">{t}</div>')

        elif para.kind == "table":
            matrix: List[List[str]] = para.payload or []
            if not matrix:
                continue
            chunks.append('<table class="pb-table">')
            header = matrix[0]
            body = matrix[1:] if len(matrix) > 1 else []
            chunks.append("<thead><tr>")
            for h in header:
                chunks.append(f"<th>{_escape_html(str(h or ''))}</th>")
            chunks.append("</tr></thead>")
            if body:
                chunks.append("<tbody>")
                for r in body:
                    chunks.append("<tr>")
                    for c in r:
                        chunks.append(f"<td>{_escape_html(str(c or ''))}</td>")
                    chunks.append("</tr>")
                chunks.append("</tbody>")
            chunks.append("</table>")
    chunks.append("</div>")
    return "\n".join(chunks)


def load_playbook_sections_for_scenario(images_cfg: Dict[str, Any], image_name: str) -> Tuple[str, List[PBSection], Path]:
    if Document is None:
        return "python-docx is not installed. Install it to render playbooks.", [], scenario_playbook_path(images_cfg, image_name)

    pb_path = scenario_playbook_path(images_cfg, image_name)
    if not pb_path.exists():
        return "No playbook uploaded for this scenario yet.", [], pb_path

    try:
        doc = Document(str(pb_path))
        sections = _docx_to_playbook_sections(doc)
        return "ok", sections, pb_path
    except Exception as e:
        return f"Failed to read playbook: {e}", [], pb_path


# =============================================================================
# Login
# =============================================================================

def _force_autofill_capture_js():
    components.html(
        """
        <script>
        (function() {
          function fire(el) {
            if (!el) return;
            el.dispatchEvent(new Event('input', {bubbles: true}));
            el.dispatchEvent(new Event('change', {bubbles: true}));
            el.dispatchEvent(new Event('keyup', {bubbles: true}));
          }
          function run() {
            const inputs = parent.document.querySelectorAll('input');
            inputs.forEach((i) => fire(i));
          }
          setTimeout(run, 250);
          setTimeout(run, 750);
          setTimeout(run, 1250);
        })();
        </script>
        """,
        height=0,
    )


def login_screen(cfg: Dict[str, Any]):
    st.markdown(
        """
        <div style='margin-top: 10vh; text-align:center; margin-bottom: 1.2rem;'>
            <div style='font-family: "Consolas", "Fira Code", monospace;
                        font-size: 2.6rem;
                        letter-spacing: 0.35rem;
                        text-transform: uppercase;
                        color: #00e0ff;
                        text-shadow: 0 0 8px rgba(56,189,248,0.45);'>
                [ AVMF ]
            </div>
            <div style='font-family: "Consolas", "Fira Code", monospace;
                        font-size: 1.1rem;
                        letter-spacing: 0.14rem;
                        margin-top: 0.35rem;
                        color: #67e8f9;'>
                Advanced Volatile Memory Forensics
            </div>
            <div style='font-size: 0.85rem; margin-top: 0.35rem; color: #9ca3af; letter-spacing: 0.08rem;'>
                Analyst Access · Memory Artifact Review
            </div>
            <hr style='margin-top:1.1rem; border: 1px solid rgba(31,41,55,0.8);'/>
        </div>
        """,
        unsafe_allow_html=True,
    )

    _force_autofill_capture_js()

    _, center, _ = st.columns([2, 3, 2])
    with center:
        with st.form("login_form", clear_on_submit=False):
            st.text_input("Username", key="login_username")
            st.text_input("Password", type="password", key="login_password")
            remember = st.checkbox("Keep me logged in (recommended)", value=True, key="login_remember")
            submitted = st.form_submit_button("Log in", use_container_width=True)

        if submitted:
            _force_autofill_capture_js()

            username_val = (st.session_state.get("login_username") or "").strip()
            password_val = st.session_state.get("login_password") or ""

            if authenticate(username_val, password_val, cfg):
                st.session_state["authenticated"] = True
                st.session_state["username"] = username_val
                st.session_state["page"] = "Dashboard"

                cm = _get_cookie_manager()
                if cm is not None and remember:
                    token = _make_token(username_val, ttl_seconds=60 * 60 * 24 * 14)
                    try:
                        expires_dt = datetime.now(timezone.utc) + timedelta(days=14)
                        cm.set("avmf_token", token, expires_at=expires_dt)
                    except TypeError:
                        cm.set("avmf_token", token)

                st.success(f"Logged in as **{username_val}**")
                st.rerun()
            else:
                st.error("Invalid username or password")


# =============================================================================
# Sidebar + navigation
# =============================================================================

NAV_PAGES = [
    ("Dashboard", "view_dashboard"),
    ("Data Upload", "upload_data"),
    ("YARA Rules", "view_yara_rules"),
    ("Reports", "view_reports"),
    ("IR Playbook", "view_playbook"),
    ("Settings", "edit_settings"),
]


def build_sidebar(uctx: UserCtx, images_cfg: Dict[str, Any]):
    with st.sidebar:
        st.markdown("### AVMF")
        st.markdown("Advanced Volatile Memory Forensics")
        st.markdown("---")

        st.markdown(f"**Logged in as**  \n{uctx.username}")
        st.markdown(f"**Role:** {uctx.role}")
        st.markdown("---")

        st.markdown("#### Navigation")
        for name, perm in NAV_PAGES:
            if can(uctx, perm):
                if st.button(name, use_container_width=True, key=f"nav_{name}"):
                    st.session_state["page"] = name

        st.markdown("---")
        st.markdown("#### Memory image")

        images_list = get_images_list(images_cfg)
        image_names = [img["name"] for img in images_list]

        if not image_names:
            st.error("No scenarios/images exist. Admin: add one in Settings → Scenario Catalog.")
            return

        current_default = st.session_state.get("selected_image", image_names[0])
        if current_default not in image_names:
            current_default = image_names[0]
        idx = image_names.index(current_default)

        selected = st.selectbox(
            "",
            image_names,
            index=idx,
            label_visibility="collapsed",
            key="sidebar_image_select",
        )
        st.session_state["selected_image"] = selected

        st.markdown("---")
        if st.button("Logout", use_container_width=True, key="logout_btn"):
            logout()


# =============================================================================
# Helpers: filters + dynamic tabs
# =============================================================================

def _find_pid_column(df: pd.DataFrame) -> Optional[str]:
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return None
    lower_map = {c.lower().strip(): c for c in df.columns}
    # common variants
    for k in ["pid", "process_id", "processid", "proc_id", "pid "]:
        if k in lower_map:
            return lower_map[k]
    # fuzzy: contains 'pid'
    for lc, orig in lower_map.items():
        if re.fullmatch(r".*\bpid\b.*", lc):
            return orig
    return None


def _find_time_column(df: pd.DataFrame) -> Optional[str]:
    if df is None or not isinstance(df, pd.DataFrame) or df.empty:
        return None
    lower_map = {c.lower().strip(): c for c in df.columns}
    candidates = [
        "time", "timestamp", "time_utc", "utc_time", "datetime", "date_time",
        "created", "created_at", "create_time", "creation_time", "createtime",
        "start_time", "started_at", "starttime",
        "end_time", "exit_time", "terminated_time", "exittime",
    ]
    for c in candidates:
        if c in lower_map:
            return lower_map[c]
    # fuzzy: contains "time"
    for lc, orig in lower_map.items():
        if "time" in lc or "timestamp" in lc or "date" in lc:
            return orig
    return None


def _apply_pid_filter(df: pd.DataFrame, pid_value: str) -> pd.DataFrame:
    if df is None or df.empty:
        return df
    if not pid_value.strip():
        return df

    pid_value = pid_value.strip()
    wanted = [p.strip() for p in pid_value.split(",") if p.strip()]
    if not wanted:
        return df

    pid_col = _find_pid_column(df)
    if not pid_col:
        return df

    s = df[pid_col].astype(str).str.strip()
    mask = False
    for w in wanted:
        mask = mask | (s == w)
    return df[mask]


def _df_is_effectively_empty(df: pd.DataFrame) -> bool:
    if df is None:
        return True
    if not isinstance(df, pd.DataFrame):
        return True
    return df.empty


def _show_df_or_na(df: pd.DataFrame, label: str):
    if df is None or df.empty:
        st.info(f"{label}: N/A for this scenario.")
        try:
            st.dataframe(df.head(0), use_container_width=True)
        except Exception:
            st.dataframe(pd.DataFrame(), use_container_width=True)
    else:
        st.dataframe(df, use_container_width=True)


def _enabled_tables_for_image(img_meta: Dict[str, Any]) -> Dict[str, bool]:
    tabs = (img_meta or {}).get("tabs", {}) or {}
    out = {}
    for k in TABLE_REGISTRY.keys():
        out[k] = bool(tabs.get(k, DEFAULT_TABLE_VISIBILITY.get(k, True)))
    return out


def _groups_for_enabled_tables(enabled: Dict[str, bool]) -> Dict[str, List[str]]:
    groups: Dict[str, List[str]] = {g: [] for g in GROUP_ORDER}
    for key, meta in TABLE_REGISTRY.items():
        if enabled.get(key, False):
            groups[meta["group"]].append(key)
    return {g: keys for g, keys in groups.items() if keys}


def _kpi_cards_html(items: List[Tuple[str, int]]) -> str:
    # items = [(label, value), ...]
    chunks = ['<div class="kpi-row">']
    for label, value in items:
        chunks.append(
            f'<div class="kpi-item"><div class="kpi-label">{_escape_html(label)}</div><div class="kpi-value">{value}</div></div>'
        )
    chunks.append("</div>")
    return "\n".join(chunks)


# =============================================================================
# YARA hits upload (multi-file) helpers
# =============================================================================

def _sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def _unique_target_name(folder: Path, original_name: str) -> str:
    base = safe_filename(Path(original_name).stem)
    ext = ".csv"
    candidate = f"{base}{ext}"
    if not (folder / candidate).exists():
        return candidate
    i = 2
    while True:
        candidate = f"{base}_{i}{ext}"
        if not (folder / candidate).exists():
            return candidate
        i += 1


def save_yara_hit_uploads(images_cfg: Dict[str, Any], image_name: str, uploads: List[Any], replace_all: bool) -> Tuple[int, int]:
    """
    Returns (saved_count, skipped_duplicates_count).
    Prevents duplicate saves on rerun via content-hash memoization.
    """
    folder = yara_hits_folder(images_cfg, image_name)
    if replace_all:
        try:
            for p in folder.glob("*.csv"):
                p.unlink()
        except Exception:
            pass

    key = image_storage_key(images_cfg, image_name)
    memo_key = f"_yara_upload_hashes__{key}"
    seen: set = st.session_state.get(memo_key, set())
    if not isinstance(seen, set):
        seen = set()

    saved = 0
    skipped = 0
    for up in uploads:
        try:
            content = up.getvalue()
        except Exception:
            try:
                content = up.read()
            except Exception:
                continue
        h = _sha256_bytes(content)
        if h in seen:
            skipped += 1
            continue

        # basic validation: can read as CSV
        try:
            tmp = pd.read_csv(BytesIO(content), nrows=5)
            if tmp is None:
                raise ValueError("invalid")
        except Exception:
            # still save raw bytes? no — avoid corrupt files
            continue

        tgt_name = _unique_target_name(folder, getattr(up, "name", "yara_hits.csv"))
        try:
            (folder / tgt_name).write_bytes(content)
            saved += 1
            seen.add(h)
        except Exception:
            pass

    st.session_state[memo_key] = seen
    return saved, skipped


# =============================================================================
# Pages
# =============================================================================

def dashboard_page(images_cfg: Dict[str, Any], image_name: str, uctx: UserCtx):
    img_meta = find_image(images_cfg, image_name) or {
        "name": image_name,
        "scenario": "",
        "os": "",
        "acquired_at": "",
        "tabs": dict(DEFAULT_TABLE_VISIBILITY),
        "hide_empty": False,
    }
    enabled = _enabled_tables_for_image(img_meta)
    hide_empty = bool(img_meta.get("hide_empty", False))

    st.title("AVMF Memory Forensics Dashboard")
    st.subheader("Image Overview")

    overview_html = f"""
    <div class="image-info-row">
        <div class="image-info-card">
            <div class="image-info-label">Image Name</div>
            <div class="image-info-value">{img_meta.get("name", "")}</div>
        </div>
        <div class="image-info-card">
            <div class="image-info-label">Scenario</div>
            <div class="image-info-value">{img_meta.get("scenario", "")}</div>
        </div>
        <div class="image-info-card">
            <div class="image-info-label">OS</div>
            <div class="image-info-value">{img_meta.get("os", "")}</div>
        </div>
        <div class="image-info-card">
            <div class="image-info-label">Acquired At</div>
            <div class="image-info-value">{img_meta.get("acquired_at", "")}</div>
        </div>
    </div>
    """
    st.markdown(overview_html, unsafe_allow_html=True)

    data = load_image_data(images_cfg, image_name)

    # Dynamic KPIs (only enabled tables)
    kpi_items: List[Tuple[str, int]] = []
    for k, meta in TABLE_REGISTRY.items():
        if not enabled.get(k, False):
            continue
        df = data.get(k, pd.DataFrame())
        kpi_items.append((meta["label"], int(len(df.index)) if isinstance(df, pd.DataFrame) else 0))
    if kpi_items:
        st.markdown(_kpi_cards_html(kpi_items), unsafe_allow_html=True)

    groups = _groups_for_enabled_tables(enabled)

    # Auto-hide empty tables
    if hide_empty:
        groups2: Dict[str, List[str]] = {}
        for g, keys in groups.items():
            kept = [k for k in keys if not _df_is_effectively_empty(data.get(k))]
            if kept:
                groups2[g] = kept
        groups = groups2

    if not groups:
        st.warning("No dashboard tabs are enabled for this scenario. Admin: enable tabs in Settings → Scenario Catalog.")
        return

    # Render top-level group tabs
    group_labels = [GROUP_LABELS[g] for g in groups.keys()]
    tab_objs = st.tabs(group_labels)

    for tab_obj, group_key in zip(tab_objs, groups.keys()):
        with tab_obj:
            keys_in_group = groups[group_key]

            pid_filter_key = f"pid_filter__{image_storage_key(images_cfg, image_name)}__{group_key}"
            pid_filter = st.text_input("Filter by PID (comma-separated, optional)", value="", key=pid_filter_key)

            def get_table_df(k: str) -> pd.DataFrame:
                df = data.get(k, pd.DataFrame())
                df = _apply_pid_filter(df, pid_filter)
                return df

            if group_key == "procnet":
                if "processes" in keys_in_group:
                    st.subheader("Processes")
                    _show_df_or_na(get_table_df("processes"), "Processes")
                if "network" in keys_in_group:
                    st.subheader("Network Connections")
                    proc_filter_key = f"net_proc_filter__{image_storage_key(images_cfg, image_name)}"
                    proc_filter = st.text_input("Filter by process name (optional)", "", key=proc_filter_key)
                    net_df = get_table_df("network")
                    if net_df is not None and (not net_df.empty) and proc_filter:
                        # try common columns
                        col = _find_col(net_df, ["process", "name", "ImageFileName", "imagefilename"])
                        if col:
                            net_df = net_df[net_df[col].astype(str).str.contains(proc_filter, case=False, na=False)]
                    _show_df_or_na(net_df, "Network Connections")

            elif group_key == "yara":
                if "yara" in keys_in_group:
                    st.subheader("YARA Hits (combined)")
                    files = list_yara_hit_files(images_cfg, image_name)
                    st.caption(f"Stored YARA hit CSVs for this scenario: {len(files)} file(s).")
                    if files:
                        with st.expander("Show file list", expanded=False):
                            for p in files:
                                st.write(f"- {p.name}")
                    _show_df_or_na(get_table_df("yara"), "YARA Hits")

            elif group_key == "persistence":
                sub_labels = []
                sub_keys = []
                if "runkeys" in keys_in_group:
                    sub_labels.append("Run Keys")
                    sub_keys.append("runkeys")
                if "runonce" in keys_in_group:
                    sub_labels.append("RunOnce")
                    sub_keys.append("runonce")

                if len(sub_keys) == 1:
                    k = sub_keys[0]
                    st.subheader(TABLE_REGISTRY[k]["label"])
                    _show_df_or_na(get_table_df(k), TABLE_REGISTRY[k]["label"])
                else:
                    sub_tabs = st.tabs(sub_labels)
                    for sub_tab, k in zip(sub_tabs, sub_keys):
                        with sub_tab:
                            _show_df_or_na(get_table_df(k), TABLE_REGISTRY[k]["label"])

            elif group_key == "cmdline":
                if "cmdline" in keys_in_group:
                    st.subheader("Command Line")
                    _show_df_or_na(get_table_df("cmdline"), "Cmdline")

            elif group_key == "sessions":
                if "sessions" in keys_in_group:
                    st.subheader("Sessions")
                    _show_df_or_na(get_table_df("sessions"), "Sessions")

            elif group_key == "logons":
                if "logons" in keys_in_group:
                    st.subheader("Logon & Credential-related Events")
                    _show_df_or_na(get_table_df("logons"), "Logon Events")

            else:
                for k in keys_in_group:
                    st.subheader(TABLE_REGISTRY[k]["label"])
                    _show_df_or_na(get_table_df(k), TABLE_REGISTRY[k]["label"])

    st.markdown("---")
    st.caption("Dashboard tables reflect CSVs you upload (admin) on the Data Upload page. Employees can view only.")


def data_upload_page(images_cfg: Dict[str, Any], image_name: str, uctx: UserCtx):
    st.title("Data Upload")
    if not can(uctx, "upload_data"):
        st.warning("Admins only: Upload per-image CSVs to drive the dashboard.")
        return

    img_meta = find_image(images_cfg, image_name) or {"name": image_name, "tabs": dict(DEFAULT_TABLE_VISIBILITY)}
    enabled = _enabled_tables_for_image(img_meta)

    st.markdown(
        "Upload per-image CSV files to drive the dashboard. "
        "If a CSV is uploaded here, it **overrides** the built-in defaults for the selected scenario/image."
    )

    # migrate legacy so nothing "disappears"
    migrate_legacy_name_based_files(images_cfg, image_name)

    paths = data_paths_for_image(images_cfg, image_name)
    st.markdown(f"**Current memory image:** `{image_name}`")
    st.markdown("---")

    def upload_section_generic(key: str):
        meta = TABLE_REGISTRY[key]
        label = meta["label"]
        path = paths[key]
        recommended_cols = meta["recommended_cols"]

        st.subheader(label)

        if path.exists():
            st.success(f"Using uploaded CSV at: `{path.name}`")
            try:
                preview_df = pd.read_csv(path, nrows=25)
                st.caption("Current file preview (first 25):")
                st.dataframe(preview_df, use_container_width=True)
            except Exception as e:
                st.error(f"Could not read existing CSV: {e}")
        else:
            st.info("Currently using built-in defaults / blank tables for this section.")

        st.caption("Recommended columns: " + ", ".join(recommended_cols) + " (extra columns are fine).")
        file = st.file_uploader("Upload CSV", type="csv", key=f"upload_{image_storage_key(images_cfg, image_name)}_{key}")
        if file is not None:
            try:
                df = pd.read_csv(file)
                path.write_text(df.to_csv(index=False), encoding="utf-8")
                st.success(f"Saved uploaded CSV to `{path.name}`. Dashboard will use it.")
                st.rerun()
            except Exception as e:
                st.error(f"Failed to save CSV: {e}")
        st.markdown("---")

    def upload_section_yara():
        meta = TABLE_REGISTRY["yara"]
        st.subheader(meta["label"])
        st.caption("Recommended columns: rule, pid, process, address, note (extra columns are fine).")

        folder = yara_hits_folder(images_cfg, image_name)
        files = list_yara_hit_files(images_cfg, image_name)

        st.info(f"YARA hit CSVs stored for this scenario: **{len(files)}** file(s). (Stored under `{folder.name}/`)")

        # Combined preview table (no dropdown)
        combined, _ = load_yara_hits_combined(images_cfg, image_name)
        if combined is not None and not combined.empty:
            st.caption("Current combined preview (first 200 rows):")
            st.dataframe(combined.head(200), use_container_width=True)
        else:
            st.caption("No YARA hit data loaded yet for this scenario.")

        st.markdown("**Upload multiple YARA hit CSVs (one per scan / rule-pack / run):**")
        replace_all = st.checkbox("Replace ALL existing YARA hit CSVs for this image", value=False,
                                  key=f"yara_replace_all__{image_storage_key(images_cfg, image_name)}")
        uploads = st.file_uploader(
            "Upload one or more CSVs",
            type="csv",
            accept_multiple_files=True,
            key=f"yara_multi_upload__{image_storage_key(images_cfg, image_name)}",
        )

        if uploads:
            saved, skipped = save_yara_hit_uploads(images_cfg, image_name, uploads, replace_all=replace_all)
            if saved > 0:
                st.success(f"Saved {saved} YARA hit CSV file(s). Dashboard will show them combined.")
                st.rerun()
            else:
                if skipped > 0:
                    st.info("Those files were already processed (duplicate upload/rerun) — nothing new saved.")
                else:
                    st.warning("No valid CSVs were saved (check file format).")

        # Deletion
        files = list_yara_hit_files(images_cfg, image_name)
        if files:
            st.markdown("**Delete stored YARA hit CSVs:**")
            options = [p.name for p in files]
            to_del = st.multiselect(
                "Select files to delete",
                options,
                default=[],
                key=f"yara_delete_select__{image_storage_key(images_cfg, image_name)}",
            )
            if st.button("Delete selected YARA CSVs", key=f"yara_delete_btn__{image_storage_key(images_cfg, image_name)}"):
                deleted = 0
                for nm in to_del:
                    p = folder / nm
                    try:
                        if p.exists():
                            p.unlink()
                            deleted += 1
                    except Exception:
                        pass
                st.success(f"Deleted {deleted} file(s).")
                st.rerun()

        st.markdown("---")

    any_enabled = False
    for k in TABLE_REGISTRY.keys():
        if not enabled.get(k, False):
            continue
        any_enabled = True
        if k == "yara":
            upload_section_yara()
        else:
            upload_section_generic(k)

    if not any_enabled:
        st.warning("No upload sections are enabled for this scenario. Enable tables in Settings → Scenario Catalog.")


def reports_page(images_cfg: Dict[str, Any], image_name: str, uctx: UserCtx):
    st.title("Forensic Reports")
    st.markdown("Generate a structured report from the current dashboard data (includes narrative + Timeline Reconstruction when timestamps exist).")

    template = load_report_template()

    if can(uctx, "edit_report_template"):
        with st.expander("Admin: Report format/template (editable)", expanded=False):
            new_t = st.text_area("Report template", value=template, height=320, key="report_template_editor")
            if st.button("Save template", type="primary", key="report_template_save"):
                save_report_template(new_t)
                st.success("Template saved.")
                template = new_t

    report_text, payload = render_report_from_dashboard(images_cfg, template, image_name)

    st.subheader("Report preview")
    st.text_area("", value=report_text, height=420, label_visibility="collapsed", key="report_preview")

    st.markdown("---")
    if can(uctx, "download_reports"):
        st.download_button(
            "⬇ Download as text (.txt)",
            data=report_text.encode("utf-8"),
            file_name=f"{image_name}_forensic_report.txt",
            mime="text/plain",
            key="dl_report_txt",
        )

        docx_bytes = build_report_docx_bytes(payload)
        if docx_bytes is not None:
            st.download_button(
                "⬇ Download Word report (.docx)",
                data=docx_bytes,
                file_name=f"{image_name}_forensic_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_report_docx",
            )
        else:
            st.info("Install python-docx to enable Word export.")


def ir_playbook_page(images_cfg: Dict[str, Any], sidebar_image_name: str, uctx: UserCtx):
    st.title("Incident Response Playbook")

    imgs = get_images_list(images_cfg)
    if not imgs:
        st.warning("No scenarios exist. Admin: add scenarios in Settings.")
        return

    names = [i["name"] for i in imgs]
    default_name = sidebar_image_name if sidebar_image_name in names else names[0]
    idx = names.index(default_name)

    # per-page selector (does not depend on sidebar)
    selected = st.selectbox("Select scenario / image (this page)", names, index=idx, key="pb_page_image_select")

    img = find_image(images_cfg, selected) or {"name": selected, "scenario": ""}
    st.caption(f"Scenario: **{img.get('scenario','')}** · Image: `{img.get('name','')}`")

    pb_path = scenario_playbook_path(images_cfg, selected)

    # Admin upload per scenario
    if can(uctx, "manage_playbook"):
        with st.expander("Admin: Scenario Playbook Management", expanded=False):
            st.markdown(f"Stored at: `{pb_path}`")
            upload = st.file_uploader(
                "Upload/replace playbook for THIS scenario (.docx)",
                type=["docx"],
                key=f"playbook_upload__{image_storage_key(images_cfg, selected)}",
            )
            if upload is not None:
                try:
                    pb_path.write_bytes(upload.read())
                    st.success("Playbook saved for this scenario.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Failed to save playbook: {e}")

            if pb_path.exists():
                if st.button("Delete this scenario playbook", key=f"pb_delete__{image_storage_key(images_cfg, selected)}"):
                    try:
                        pb_path.unlink()
                        st.success("Deleted.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Failed: {e}")

    if not pb_path.exists():
        st.info("No playbook uploaded for this scenario yet.")
        return

    if can(uctx, "download_playbook"):
        st.download_button(
            "⬇ Download Active Playbook (.docx)",
            data=pb_path.read_bytes(),
            file_name=pb_path.name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_playbook__{image_storage_key(images_cfg, selected)}",
        )

    if not can(uctx, "view_playbook"):
        st.warning("You do not have permission to view this playbook.")
        return

    status, sections, _ = load_playbook_sections_for_scenario(images_cfg, selected)
    if not sections:
        st.error(status)
        return

    top_sections = [s for s in sections if s.level == 1]
    options = ["Full Playbook"] + [s.title for s in top_sections]

    c1, c2 = st.columns([2, 2])
    with c1:
        search_q = st.text_input("Search inside playbook", "", key=f"pb_search__{image_storage_key(images_cfg, selected)}")
    with c2:
        selection = st.selectbox(
            "Jump to (filters content)",
            options,
            index=0,
            key=f"pb_jump__{image_storage_key(images_cfg, selected)}",
        )

    def section_to_paras(sec: PBSection) -> List[PBPara]:
        paras: List[PBPara] = []
        paras.append(PBPara("h1" if sec.level == 1 else ("h2" if sec.level == 2 else "h3"), sec.title))
        paras.extend(sec.paras)
        return paras

    if selection == "Full Playbook":
        filtered_paras: List[PBPara] = []
        for s in sections:
            filtered_paras.extend(section_to_paras(s))
    else:
        start_i = next((i for i, s in enumerate(sections) if s.level == 1 and s.title == selection), None)
        end_i = None
        if start_i is not None:
            for j in range(start_i + 1, len(sections)):
                if sections[j].level == 1:
                    end_i = j
                    break
            if end_i is None:
                end_i = len(sections)
            filtered_paras = []
            for s in sections[start_i:end_i]:
                filtered_paras.extend(section_to_paras(s))
        else:
            filtered_paras = []
            for s in sections:
                filtered_paras.extend(section_to_paras(s))

    if search_q.strip():
        q = search_q.strip().lower()

        def matches(p: PBPara) -> bool:
            if p.kind == "table":
                matrix = p.payload or []
                for r in matrix:
                    for c in r:
                        if q in str(c or "").lower():
                            return True
                return False
            return q in str(p.payload or "").lower()

        kept: List[PBPara] = []
        for p in filtered_paras:
            if p.kind in ("h1", "h2", "h3"):
                kept.append(p)
            elif matches(p):
                kept.append(p)
        filtered_paras = kept

    st.markdown(_render_playbook_html(filtered_paras), unsafe_allow_html=True)


# =============================================================================
# Settings (scenario catalog + RBAC) + rename scenario name FIX
# =============================================================================

def settings_page(cfg: Dict[str, Any], images_cfg: Dict[str, Any], uctx: UserCtx):
    st.title("Settings")
    if not can(uctx, "edit_settings"):
        st.warning("Admins only.")
        return

    st.subheader("Appearance")
    st.checkbox("Wide mode", value=st.session_state.get("ui_wide_mode", True), key="ui_wide_mode")
    theme_options = ["Use system setting", "Light", "Dark"]
    theme_map = {"Use system setting": "system", "Light": "light", "Dark": "dark"}
    reverse_map = {v: k for k, v in theme_map.items()}
    current_label = reverse_map.get(st.session_state.get("app_theme", "system"), "Use system setting")
    choice = st.selectbox("Choose app theme", theme_options, index=theme_options.index(current_label), key="settings_theme")
    st.session_state["app_theme"] = theme_map[choice]

    st.markdown("---")
    st.subheader("Scenario Catalog (Memory Images)")

    if not can(uctx, "manage_images"):
        st.info("You do not have permission to manage scenarios/images.")
    else:
        imgs = get_images_list(images_cfg)
        table_rows = []
        for img in imgs:
            _normalize_image_tabs(img)
            t = img.get("tabs", {})
            table_rows.append({
                "id": img.get("id"),
                "name": img.get("name"),
                "scenario": img.get("scenario"),
                "os": img.get("os"),
                "acquired_at": img.get("acquired_at"),
                "pipeline": img.get("pipeline", "generic"),
                "hide_empty": img.get("hide_empty", False),
                "tabs_enabled_count": sum(1 for k in TABLE_REGISTRY.keys() if bool(t.get(k, True))),
                "has_playbook": scenario_playbook_path(images_cfg, img.get("name", "")).exists(),
                "yara_hit_files": len(list_yara_hit_files(images_cfg, img.get("name", ""))),
            })
        st.dataframe(pd.DataFrame(table_rows), use_container_width=True)

        st.markdown("### Add new scenario/image")
        col1, col2 = st.columns([1, 1])
        with col1:
            new_name = st.text_input("Memory dump filename / label (unique)", value="", key="img_add_name")
            new_scenario = st.text_input("Scenario name", value="", key="img_add_scenario")
            new_os = st.text_input("Operating OS", value="Windows 11", key="img_add_os")
        with col2:
            new_acq = st.text_input("Acquired At", value="", key="img_add_acquired")
            new_pipeline = st.selectbox("Pipeline", ["generic", "locky"], index=0, key="img_add_pipeline")
            st.caption("generic = blank tables until CSVs. locky = uses locky_* normalized if present.")

        if st.button("Add scenario", type="primary", key="btn_add_image"):
            nn = (new_name or "").strip()
            if not nn:
                st.error("Memory dump filename/label is required.")
            elif find_image(images_cfg, nn) is not None:
                st.error("That scenario name already exists.")
            else:
                images_cfg["images"].append({
                    "id": next_image_id(images_cfg),
                    "name": nn,
                    "scenario": (new_scenario or "").strip(),
                    "os": (new_os or "").strip(),
                    "acquired_at": (new_acq or "").strip(),
                    "pipeline": new_pipeline,
                    "tabs": dict(DEFAULT_TABLE_VISIBILITY),
                    "hide_empty": False,
                })
                save_images_config(images_cfg)
                st.success("Scenario added.")
                st.session_state["selected_image"] = nn
                st.rerun()

        st.markdown("### Edit / delete existing scenario/image")
        if imgs:
            target_name = st.selectbox("Select image", [i["name"] for i in imgs], key="img_edit_select")
            tgt = find_image(images_cfg, target_name)
            if tgt:
                _normalize_image_tabs(tgt)

                old_name = str(tgt.get("name", "") or "")
                k_name = f"img_edit_name__{int(tgt.get('id',0))}"
                k_scn = f"img_edit_scenario__{int(tgt.get('id',0))}"
                k_os = f"img_edit_os__{int(tgt.get('id',0))}"
                k_acq = f"img_edit_acquired__{int(tgt.get('id',0))}"
                k_pipe = f"img_edit_pipeline__{int(tgt.get('id',0))}"
                k_cleanup = f"img_delete_cleanup__{int(tgt.get('id',0))}"
                k_hide_empty = f"img_hide_empty__{int(tgt.get('id',0))}"
                k_rename_legacy = f"img_rename_legacy__{int(tgt.get('id',0))}"

                st.subheader("Scenario Metadata")
                e1, e2 = st.columns([1, 1])
                with e1:
                    edit_name = st.text_input("Memory dump filename / label (rename scenario)", value=old_name, key=k_name)
                    edit_scenario = st.text_input("Scenario", value=tgt.get("scenario", ""), key=k_scn)
                with e2:
                    edit_os = st.text_input("OS", value=tgt.get("os", ""), key=k_os)
                    edit_acq = st.text_input("Acquired At", value=tgt.get("acquired_at", ""), key=k_acq)

                edit_pipeline = st.selectbox(
                    "Pipeline",
                    ["generic", "locky"],
                    index=(0 if tgt.get("pipeline", "generic") == "generic" else 1),
                    key=k_pipe,
                )

                rename_legacy = st.checkbox(
                    "Also rename legacy name-based CSV + YARA rule file(s) on disk (optional)",
                    value=True,
                    key=k_rename_legacy,
                )

                st.markdown("---")
                st.subheader("Dashboard Tabs for This Scenario")
                st.caption("Turn on/off tabs and subtabs (tables). Data Upload page will match this automatically.")

                tabs_cfg = dict(tgt.get("tabs", {}))
                toggles: Dict[str, bool] = {}
                cols = st.columns(2)
                items = list(TABLE_REGISTRY.keys())
                for i, key in enumerate(items):
                    col = cols[i % 2]
                    with col:
                        toggles[key] = st.checkbox(
                            f"{TABLE_REGISTRY[key]['icon']} {TABLE_REGISTRY[key]['label']}",
                            value=bool(tabs_cfg.get(key, True)),
                            key=f"tab_toggle__{int(tgt.get('id',0))}__{key}",
                        )

                hide_empty_val = st.checkbox(
                    "Auto-hide empty tables/tabs for this scenario (instead of showing N/A)",
                    value=bool(tgt.get("hide_empty", False)),
                    key=k_hide_empty,
                )

                cR1, cR2 = st.columns([1, 1])
                with cR1:
                    if st.button("Reset tabs to default (enable all)", key=f"btn_reset_tabs__{int(tgt.get('id',0))}"):
                        tgt["tabs"] = dict(DEFAULT_TABLE_VISIBILITY)
                        tgt["hide_empty"] = False
                        save_images_config(images_cfg)
                        st.success("Reset done.")
                        st.rerun()
                with cR2:
                    if st.button("Disable all tabs", key=f"btn_disable_all_tabs__{int(tgt.get('id',0))}"):
                        tgt["tabs"] = {k: False for k in TABLE_REGISTRY.keys()}
                        tgt["hide_empty"] = True
                        save_images_config(images_cfg)
                        st.success("All tabs disabled.")
                        st.rerun()

                st.markdown("---")
                cA, cB = st.columns([1, 1])
                with cA:
                    if st.button("Save scenario changes", type="primary", key=f"btn_save_image__{int(tgt.get('id',0))}"):
                        nn = (edit_name or "").strip()
                        if not nn:
                            st.error("Scenario name cannot be empty.")
                            st.stop()
                        if nn != old_name and find_image(images_cfg, nn) is not None:
                            st.error("That scenario name already exists. Choose a unique name.")
                            st.stop()

                        # rename legacy files if requested (stable ID-based storage does NOT depend on name)
                        if rename_legacy and nn != old_name:
                            try:
                                old_legacy = legacy_data_paths_for_image(old_name)
                                new_legacy = legacy_data_paths_for_image(nn)
                                for tk, op in old_legacy.items():
                                    np = new_legacy.get(tk)
                                    if op.exists() and np and not np.exists():
                                        try:
                                            op.rename(np)
                                        except Exception:
                                            pass
                            except Exception:
                                pass
                            try:
                                old_yar = _yara_file_for_image(old_name)
                                new_yar = _yara_file_for_image(nn)
                                if old_yar.exists() and not new_yar.exists():
                                    old_yar.rename(new_yar)
                            except Exception:
                                pass

                        tgt["name"] = nn
                        tgt["scenario"] = edit_scenario
                        tgt["os"] = edit_os
                        tgt["acquired_at"] = edit_acq
                        tgt["pipeline"] = edit_pipeline
                        tgt["tabs"] = toggles
                        tgt["hide_empty"] = hide_empty_val
                        save_images_config(images_cfg)

                        st.session_state["selected_image"] = nn  # prevents "revert"
                        st.success("Saved.")
                        st.rerun()

                with cB:
                    cleanup = st.checkbox("Also delete this image's CSVs + YARA hits + scenario playbook", value=False, key=k_cleanup)
                    if st.button("Delete scenario", type="secondary", key=f"btn_delete_image__{int(tgt.get('id',0))}"):
                        target_id = int(tgt.get("id", 0))
                        target_name_now = str(tgt.get("name", ""))

                        images_cfg["images"] = [i for i in images_cfg["images"] if int(i.get("id", 0)) != target_id]
                        save_images_config(images_cfg)

                        if cleanup:
                            # delete stable ID-based CSVs
                            key = image_storage_key(images_cfg, target_name_now)  # uses fallback if removed; compute directly
                            key = f"img_{target_id}"
                            for tkey, meta in TABLE_REGISTRY.items():
                                p = DATA_DIR / f"{key}_{meta['filename_suffix']}.csv"
                                try:
                                    if p.exists():
                                        p.unlink()
                                except Exception:
                                    pass
                            # delete YARA hits folder
                            try:
                                f = YARA_HITS_DIR / key
                                if f.exists():
                                    shutil.rmtree(f, ignore_errors=True)
                            except Exception:
                                pass
                            # delete playbook
                            try:
                                pb = PLAYBOOKS_DIR / f"{key}.docx"
                                if pb.exists():
                                    pb.unlink()
                            except Exception:
                                pass
                            # delete YARA rules file name-based (optional)
                            try:
                                yf = _yara_file_for_image(target_name_now)
                                if yf.exists():
                                    yf.unlink()
                            except Exception:
                                pass

                        remaining = get_images_list(images_cfg)
                        st.session_state["selected_image"] = remaining[0]["name"] if remaining else None
                        st.success("Deleted.")
                        st.rerun()

    st.markdown("---")
    st.subheader("User Management (RBAC)")

    if not can(uctx, "manage_users"):
        st.warning("You do not have permission to manage users.")
        return

    users = cfg.get("users", {})

    view_rows = []
    for uname, u in users.items():
        perms = u.get("perms", {})
        view_rows.append({
            "username": uname,
            "role": u.get("role", "employee"),
            "view_dashboard": perms.get("view_dashboard", False),
            "view_reports": perms.get("view_reports", False),
            "download_reports": perms.get("download_reports", False),
            "view_playbook": perms.get("view_playbook", False),
            "download_playbook": perms.get("download_playbook", False),
            "upload_data": perms.get("upload_data", False),
            "view_yara_rules": perms.get("view_yara_rules", False),
            "edit_report_template": perms.get("edit_report_template", False),
            "manage_images": perms.get("manage_images", False),
            "manage_users": perms.get("manage_users", False),
        })
    st.dataframe(pd.DataFrame(view_rows), use_container_width=True)

    st.markdown("---")
    colA, colB = st.columns([1, 1])

    with colA:
        st.markdown("#### Add user")
        new_user = st.text_input("New username", key="new_user_name")
        new_pass = st.text_input("New password", type="password", key="new_user_pass")
        new_role = st.selectbox("Role", ["employee", "admin"], index=0, key="new_user_role")

        if st.button("Add user", type="primary", key="btn_add_user"):
            if not new_user or not new_pass:
                st.error("Username and password required.")
            elif new_user in users:
                st.error("User already exists.")
            else:
                salt = base64.urlsafe_b64encode(os.urandom(6)).decode("utf-8")
                if new_role == "admin":
                    perms = {
                        "view_dashboard": True, "view_reports": True, "download_reports": True,
                        "edit_report_template": True,
                        "view_playbook": True, "download_playbook": True, "manage_playbook": True,
                        "upload_data": True,
                        "view_yara_rules": True, "edit_yara_rules": True,
                        "manage_users": True, "edit_settings": True, "manage_images": True,
                    }
                else:
                    perms = {
                        "view_dashboard": True,
                        "view_reports": True, "download_reports": True,
                        "edit_report_template": False,
                        "view_playbook": True, "download_playbook": True, "manage_playbook": False,
                        "upload_data": False,
                        "view_yara_rules": False, "edit_yara_rules": False,
                        "manage_users": False, "edit_settings": False, "manage_images": False,
                    }

                users[new_user] = {"role": new_role, "salt": salt, "password_hash": _hash_password(new_pass, salt), "perms": perms}
                cfg["users"] = users
                save_users_config(cfg)
                st.success("User added.")
                st.rerun()

    with colB:
        st.markdown("#### Edit / remove user")
        if not users:
            st.info("No users found.")
        else:
            target = st.selectbox("Select user", list(users.keys()), key="edit_user_select")
            tu = users.get(target, {})

            role = st.selectbox("Role", ["employee", "admin"], index=(0 if tu.get("role") == "employee" else 1), key="edit_user_role")
            perms = dict(tu.get("perms", {}))

            st.markdown("**Permissions**")
            perms["view_dashboard"] = st.checkbox("View Dashboard", value=perms.get("view_dashboard", False), key="perm_view_dashboard")
            perms["view_reports"] = st.checkbox("View Reports", value=perms.get("view_reports", False), key="perm_view_reports")
            perms["download_reports"] = st.checkbox("Download Reports", value=perms.get("download_reports", False), key="perm_download_reports")
            perms["edit_report_template"] = st.checkbox("Edit Report Template", value=perms.get("edit_report_template", False), key="perm_edit_report_template")

            perms["view_playbook"] = st.checkbox("View Playbook", value=perms.get("view_playbook", False), key="perm_view_playbook")
            perms["download_playbook"] = st.checkbox("Download Playbook", value=perms.get("download_playbook", False), key="perm_download_playbook")
            perms["manage_playbook"] = st.checkbox("Manage Playbook File", value=perms.get("manage_playbook", False), key="perm_manage_playbook")

            perms["upload_data"] = st.checkbox("Upload Data", value=perms.get("upload_data", False), key="perm_upload_data")
            perms["view_yara_rules"] = st.checkbox("View YARA Rules", value=perms.get("view_yara_rules", False), key="perm_view_yara")
            perms["edit_yara_rules"] = st.checkbox("Edit YARA Rules", value=perms.get("edit_yara_rules", False), key="perm_edit_yara")

            perms["manage_images"] = st.checkbox("Manage Scenarios (Images)", value=perms.get("manage_images", False), key="perm_manage_images")
            perms["manage_users"] = st.checkbox("Manage Users", value=perms.get("manage_users", False), key="perm_manage_users")
            perms["edit_settings"] = st.checkbox("Edit Settings", value=perms.get("edit_settings", False), key="perm_edit_settings")

            if st.button("Save changes", type="primary", key="btn_save_user_changes"):
                users[target]["role"] = role
                users[target]["perms"] = perms
                cfg["users"] = users
                save_users_config(cfg)
                st.success("Saved.")
                st.rerun()

            st.markdown("---")
            if target == uctx.username:
                st.info("You can't delete your own active account here.")
            else:
                if st.button("Delete user", type="secondary", key="btn_delete_user"):
                    users.pop(target, None)
                    cfg["users"] = users
                    save_users_config(cfg)
                    st.success("Deleted.")
                    st.rerun()


# =============================================================================
# Main routing
# =============================================================================

def main():
    if "page" not in st.session_state:
        st.session_state["page"] = "Dashboard"
    if "ui_wide_mode" not in st.session_state:
        st.session_state["ui_wide_mode"] = True
    if "app_theme" not in st.session_state:
        st.session_state["app_theme"] = "system"
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
        st.session_state["username"] = None

    apply_theme(theme=st.session_state["app_theme"], wide_mode=st.session_state["ui_wide_mode"])

    # Cookie manager hydration
    cm = _get_cookie_manager()
    if cm is not None and not st.session_state.get("_cookie_hydrated_once"):
        st.session_state["_cookie_hydrated_once"] = True
        try:
            _ = cm.get("avmf_token")
        except Exception:
            pass
        st.rerun()

    cfg = load_users_config()
    images_cfg = load_images_config()

    uctx = current_user(cfg)
    if not uctx:
        login_screen(cfg)
        return

    images_list = get_images_list(images_cfg)
    if images_list and "selected_image" not in st.session_state:
        st.session_state["selected_image"] = images_list[0]["name"]
    elif images_list and st.session_state.get("selected_image") not in [i["name"] for i in images_list]:
        st.session_state["selected_image"] = images_list[0]["name"]

    build_sidebar(uctx, images_cfg)

    page = st.session_state.get("page", "Dashboard")
    image_name = st.session_state.get("selected_image", images_list[0]["name"] if images_list else "")

    perm_for_page = dict(NAV_PAGES).get(page)
    if perm_for_page and not can(uctx, perm_for_page):
        st.session_state["page"] = "Dashboard"
        st.warning("Access denied. Redirected to Dashboard.")
        st.rerun()

    if page == "Dashboard":
        dashboard_page(images_cfg, image_name, uctx)
    elif page == "Data Upload":
        data_upload_page(images_cfg, image_name, uctx)
    elif page == "YARA Rules":
        yara_rules_page(image_name, uctx)
    elif page == "Reports":
        reports_page(images_cfg, image_name, uctx)
    elif page == "IR Playbook":
        ir_playbook_page(images_cfg, image_name, uctx)
    elif page == "Settings":
        settings_page(cfg, images_cfg, uctx)
    else:
        st.session_state["page"] = "Dashboard"
        st.rerun()


if __name__ == "__main__":
    main()
